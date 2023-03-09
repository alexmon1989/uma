from apps.search.dataclasses import InidCode
from apps.search.templatetags.search_extras import get_person_name, get_person_country
from apps.bulletin.services import bulletin_get_number_441_code
from templatetags.uma_extras import list_of_dicts_unique

from typing import List
from abc import ABC, abstractmethod
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.image.exceptions import UnrecognizedImageError

from django.conf import settings


class ReportItem(ABC):
    """Интерфейс объекта пром. собств., который попадает в отчёт."""
    application_data: dict
    ipc_fields: List[InidCode]  # Список полей для отображения с флагом доступно ли поле для отображения

    def _get_inid(self, obj_type_id: int, code: str, obj_state: int = 1) -> InidCode | None:
        """Возвращает признак необходимости отображения поля."""
        try:
            return next(
                filter(
                    lambda x: x.obj_type_id == obj_type_id and x.code == code and x.obj_state == obj_state,
                    self.ipc_fields
                )
            )
        except StopIteration:
            return None


class ReportItemDocx(ReportItem):

    @abstractmethod
    def write(self, document: Document) -> Paragraph:
        """Возвращает абзац текста с информацией об объекте пром. собств."""
        raise NotImplemented

    def __init__(self, application_data: dict, ipc_fields: List[InidCode], lang_code: str = 'ua'):
        self.application_data = application_data
        self.ipc_fields = ipc_fields
        self.lang_code = lang_code


class ReportItemDocxTM(ReportItemDocx):
    """Торговая марка."""
    _paragraph: Paragraph
    document: Document
    obj_type_id = 4

    def _write_540(self) -> None:
        """Добавляет изображение в документ."""
        inid = self._get_inid(self.obj_type_id, '540', self.application_data['search_data']['obj_state'])
        mark_image_filename = self.application_data['TradeMark']['TrademarkDetails'].get(
            'MarkImageDetails', {}
        ).get(
            'MarkImage', {}
        ).get(
            'MarkImageFilename'
        )
        if inid and inid.visible and mark_image_filename:
            path = Path(self.application_data['Document']['filesPath'].replace('\\', '/'))
            parts_len = len(path.parts)
            # Путь к изображению на диске
            mark_image_filepath = Path(settings.MEDIA_ROOT) / path.parts[parts_len - 3]
            mark_image_filepath /= mark_image_filepath / path.parts[parts_len - 2]
            mark_image_filepath /= path.parts[parts_len - 1]
            mark_image_filepath /= mark_image_filename
            try:
                run = self._paragraph.add_run()
                run.add_picture(str(mark_image_filepath), width=Inches(2.5))
                self._paragraph.add_run('\r')
            except (UnrecognizedImageError, ZeroDivisionError):
                pass

    def _write_111(self) -> None:
        """Записывает номер свидетельства в документ."""
        inid = self._get_inid(self.obj_type_id, '111', self.application_data['search_data']['obj_state'])
        registration_number = self.application_data['TradeMark']['TrademarkDetails'].get('RegistrationNumber')
        if inid and inid.visible and registration_number:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(registration_number).bold = True
            self._paragraph.add_run('\r')

    def _write_141(self) -> None:
        """Записывает в документ данные об ИНИД (141) Дата закінчення строку дії реєстрації знака."""
        inid = self._get_inid(self.obj_type_id, '141', self.application_data['search_data']['obj_state'])
        termination_date = self.application_data['TradeMark']['TrademarkDetails'].get('TerminationDate')
        if inid and inid.visible and termination_date:
            termination_date = datetime.strptime(termination_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(termination_date).bold = True
            self._paragraph.add_run('\r')

    def _write_151(self) -> None:
        """Записывает в документ данные об ИНИД (151) Дата реєстрації."""
        inid = self._get_inid(self.obj_type_id, '151', self.application_data['search_data']['obj_state'])
        registration_date = self.application_data['TradeMark']['TrademarkDetails'].get('RegistrationDate')
        if inid and inid.visible and registration_date:
            registration_date = datetime.strptime(registration_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(registration_date).bold = True
            self._paragraph.add_run('\r')

    def _write_181(self) -> None:
        """Записывает в документ данные об ИНИД (181) Очікувана дата закінчення строку дії реєстрації."""
        inid = self._get_inid(self.obj_type_id, '181', self.application_data['search_data']['obj_state'])
        expiry_date = self.application_data['TradeMark']['TrademarkDetails'].get('ExpiryDate')
        if inid and inid.visible and expiry_date:
            expiry_date = datetime.strptime(expiry_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(expiry_date).bold = True
            self._paragraph.add_run('\r')

    def _write_186(self) -> None:
        """Записывает в документ данные об ИНИД (186) Очікувана дата продовження строку дії реєстрації."""
        inid = self._get_inid(self.obj_type_id, '186', self.application_data['search_data']['obj_state'])
        prolongation_expiry_date = self.application_data['TradeMark']['TrademarkDetails'].get('ProlonagationExpiryDate')
        if inid and inid.visible and prolongation_expiry_date:
            prolongation_expiry_date = datetime.strptime(prolongation_expiry_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(prolongation_expiry_date).bold = True
            self._paragraph.add_run('\r')

    def _write_210(self) -> None:
        """Записывает в документ данные об ИНИД (210) Порядковий номер заявки."""
        inid = self._get_inid(self.obj_type_id, '210', self.application_data['search_data']['obj_state'])
        app_number = self.application_data['TradeMark']['TrademarkDetails'].get('ApplicationNumber')
        if inid and inid.visible and app_number:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(app_number).bold = True
            self._paragraph.add_run('\r')

    def _write_220(self) -> None:
        """Записывает в документ данные об ИНИД (220) Дата подання заявки."""
        inid = self._get_inid(self.obj_type_id, '220', self.application_data['search_data']['obj_state'])
        app_date = self.application_data['TradeMark']['TrademarkDetails'].get('ApplicationDate')
        if inid and inid.visible and app_date:
            app_date = datetime.strptime(app_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(app_date).bold = True
            self._paragraph.add_run('\r')

    def _write_app_input_date(self) -> None:
        """Записывает в документ данные о поле "Дата надходження матеріалів заявки до НОІВ"."""
        inid = self._get_inid(self.obj_type_id, '221', self.application_data['search_data']['obj_state'])
        app_input_date = self.application_data['TradeMark']['TrademarkDetails'].get('app_input_date')
        if inid and inid.visible and app_input_date:
            app_input_date = app_input_date[:10]
            app_date = datetime.strptime(app_input_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(app_date).bold = True
            self._paragraph.add_run('\r')

    def _write_531(self) -> None:
        """Записывает в документ данные об ИНИД (531) Віденська класифікація."""
        inid = self._get_inid(self.obj_type_id, '531', self.application_data['search_data']['obj_state'])
        vienna_classes = self.application_data['TradeMark']['TrademarkDetails'].get(
            'MarkImageDetails', {}
        ).get(
            'MarkImage', {}
        ).get(
            'MarkImageCategory', {}
        ).get(
            'CategoryCodeDetails', {}
        ).get(
            'CategoryCode'
        )
        if inid and inid.visible and vienna_classes:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for vienna_class in vienna_classes:
                self._paragraph.add_run(f"\r{vienna_class}").bold = True
            self._paragraph.add_run('\r')

    def _write_300(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (300) Дані щодо пріоритету відповідно до Паризької конвенції та інші дані,
        пов'язані зі старшинством або реєстрацією знака у країні походження.
        TODO: Добавить приоритет по классам.
        """
        inid = self._get_inid(self.obj_type_id, '300', self.application_data['search_data']['obj_state'])
        priority = self.application_data['TradeMark']['TrademarkDetails'].get(
            'PriorityDetails', {}
        ).get(
            'Priority'
        )
        if inid and inid.visible and priority:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for item in priority:
                self._paragraph.add_run('\r')
                if 'PriorityNumber' in item:
                    self._paragraph.add_run(f"{item['PriorityNumber']}; ").bold = True
                if 'PriorityDate' in item:
                    self._paragraph.add_run(f"{item['PriorityDate']}; ").bold = True
                if 'PriorityCountryCode' in item:
                    self._paragraph.add_run(f"{item['PriorityCountryCode']}").bold = True
            self._paragraph.add_run('\r')

    def _write_731(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (731) Ім'я та адреса заявника
        """
        inid = self._get_inid(self.obj_type_id, '731', self.application_data['search_data']['obj_state'])
        applicants = self.application_data['TradeMark']['TrademarkDetails'].get(
            'ApplicantDetails', {}
        ).get(
            'Applicant'
        )
        if inid and inid.visible and applicants:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for applicant in applicants:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        applicant['ApplicantAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                if self.application_data['search_data']['obj_state'] == 2:
                    self._paragraph.add_run('\r')
                    try:
                        self._paragraph.add_run(
                            applicant['ApplicantAddressBook']['FormattedNameAddress']['Address']['FreeFormatAddress'][
                                'FreeFormatAddressLine']
                        )
                    except KeyError:
                        pass

                    try:
                        country = applicant['ApplicantAddressBook']['FormattedNameAddress']['Address'][
                            'AddressCountryCode']
                        self._paragraph.add_run(
                            f" ({country})"
                        )
                    except KeyError:
                        pass
            self._paragraph.add_run('\r')

    def _write_732(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (732) Ім'я та адреса володільця реєстрації
        """
        inid = self._get_inid(self.obj_type_id, '732', self.application_data['search_data']['obj_state'])
        holders = self.application_data['TradeMark']['TrademarkDetails'].get(
            'HolderDetails', {}
        ).get(
            'Holder'
        )
        if inid and inid.visible and holders:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for holder in holders:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        holder['HolderAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        holder['HolderAddressBook']['FormattedNameAddress']['Address']['FreeFormatAddress'][
                            'FreeFormatAddressLine']
                    )
                except KeyError:
                    pass

                try:
                    country = holder['HolderAddressBook']['FormattedNameAddress']['Address']['AddressCountryCode']
                    self._paragraph.add_run(
                        f" ({country})"
                    )
                except KeyError:
                    pass
            self._paragraph.add_run('\r')

    def _write_740(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (740) Ім'я та адреса володільця реєстрації
        """
        inid = self._get_inid(self.obj_type_id, '740', self.application_data['search_data']['obj_state'])
        representatives = self.application_data['TradeMark']['TrademarkDetails'].get(
            'RepresentativeDetails', {}
        ).get(
            'Representative'
        )
        if inid and inid.visible and representatives:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for representative in representatives:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        representative['RepresentativeAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        representative['RepresentativeAddressBook']['FormattedNameAddress']['Address'][
                            'FreeFormatAddress']['FreeFormatAddressLine']
                    )
                except KeyError:
                    pass

                try:
                    country = representative['RepresentativeAddressBook']['FormattedNameAddress']['Address'][
                        'AddressCountryCode']
                    self._paragraph.add_run(
                        f" ({country})"
                    )
                except KeyError:
                    pass
            self._paragraph.add_run('\r')

    def _write_750(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (750) Адреса для листування
        """
        inid = self._get_inid(self.obj_type_id, '750', self.application_data['search_data']['obj_state'])
        correspondence = self.application_data['TradeMark']['TrademarkDetails'].get(
            'CorrespondenceAddress', {}
        ).get(
            'CorrespondenceAddressBook'
        )
        if inid and inid.visible and correspondence:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            self._paragraph.add_run('\r')
            try:
                self._paragraph.add_run(
                    correspondence['Name']['FreeFormatNameLine']
                ).bold = True
            except KeyError:
                pass

            self._paragraph.add_run('\r')
            try:
                self._paragraph.add_run(
                    correspondence['Address']['FreeFormatAddressLine']
                )
            except KeyError:
                pass

            try:
                country = correspondence['Address']['AddressCountryCode']
                self._paragraph.add_run(
                    f" ({country})"
                )
            except KeyError:
                pass
            self._paragraph.add_run('\r')

    def _write_591(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (591) Інформація щодо заявлених кольорів
        """
        inid = self._get_inid(self.obj_type_id, '591', self.application_data['search_data']['obj_state'])
        colors = self.application_data['TradeMark']['TrademarkDetails'].get(
            'MarkImageDetails', {}
        ).get(
            'MarkImage', {}
        ).get(
            'MarkImageColourClaimedText'
        )
        if inid and inid.visible and colors:
            self._paragraph.add_run(f"({inid.code})").bold = True
            colors_str = ', '.join([x['#text'] for x in colors])
            self._paragraph.add_run(f"\t{inid.title}: {colors_str}")
            self._paragraph.add_run('\r')

    def _write_511(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (511) Індекси Ніццької класифікації
        """
        inid = self._get_inid(self.obj_type_id, '511', self.application_data['search_data']['obj_state'])
        if 'GoodsServicesDetails' in self.application_data['TradeMark']['TrademarkDetails'] and \
                self.application_data['TradeMark']['TrademarkDetails']['GoodsServicesDetails'] is not None:
            goods = self.application_data['TradeMark']['TrademarkDetails']['GoodsServicesDetails'].get(
                'GoodsServices', {}
            ).get(
                'ClassDescriptionDetails', {}
            ).get(
                'ClassDescription'
            )
            if inid and inid.visible and goods:
                self._paragraph.add_run(f"({inid.code})").bold = True
                goods_classes_str = ', '.join([str(x['ClassNumber']) for x in goods])
                self._paragraph.add_run(f"\t{inid.title}: ")
                self._paragraph.add_run(goods_classes_str).bold = True
                for item in goods:
                    self._paragraph.add_run('\r')
                    self._paragraph.add_run(f"Кл. {item['ClassNumber']}:\t").bold = True
                    terms = item['ClassificationTermDetails']['ClassificationTerm']
                    values_str = '; '.join([x['ClassificationTermText'] for x in terms])
                    self._paragraph.add_run(values_str)
                self._paragraph.add_run('\r')

    def _write_441(self) -> None:
        """Записывает в документ данные об ИНИД (441) Дата публікації відомостей про заявку та номер бюлетня."""
        inid = self._get_inid(self.obj_type_id, '441', self.application_data['search_data']['obj_state'])
        code_441 = self.application_data['TradeMark']['TrademarkDetails'].get('Code_441')
        if inid and inid.visible and code_441:
            bul_number = self.application_data['TradeMark']['TrademarkDetails'].get('Code_441_BulNumber')
            if not bul_number:
                bul_number = str(bulletin_get_number_441_code(code_441))

            code_441 = datetime.strptime(code_441, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(code_441).bold = True
            if self.lang_code == 'ua':
                self._paragraph.add_run(', бюл. №').bold = True
            else:
                self._paragraph.add_run(', bul. №').bold = True
            self._paragraph.add_run(str(bul_number)).bold = True
            self._paragraph.add_run()
            self._paragraph.add_run('\r')

    def _write_450(self) -> None:
        """Записывает в документ данные об ИНИД (450) Дата публікації відомостей про видачу свідоцтва."""
        inid = self._get_inid(self.obj_type_id, '450', self.application_data['search_data']['obj_state'])
        publication_details = self.application_data['TradeMark']['TrademarkDetails'].get('PublicationDetails')
        if inid and inid.visible and publication_details and len(publication_details) > 0:
            bul_number = publication_details[0]['PublicationIdentifier']
            publication_date = datetime.strptime(
                publication_details[0]['PublicationDate'],
                '%Y-%m-%d'
            ).strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(publication_date).bold = True
            if self.lang_code == 'ua':
                self._paragraph.add_run(', бюл. № ').bold = True
            else:
                self._paragraph.add_run(', bul. № ').bold = True
            self._paragraph.add_run(bul_number).bold = True
            self._paragraph.add_run()
            self._paragraph.add_run('\r')

    def write(self, document: Document) -> Paragraph:
        """Записывает информацию о ТМ в абзац и возвращает его."""
        self.document = document
        self._paragraph = self.document.add_paragraph('')

        self._write_540()
        self._write_app_input_date()
        self._write_111()
        self._write_151()
        self._write_141()
        self._write_181()
        self._write_186()
        self._write_210()
        self._write_220()
        self._write_531()
        self._write_300()
        self._write_441()
        self._write_450()
        self._write_731()
        self._write_732()
        self._write_740()
        self._write_750()
        self._write_591()
        self._write_511()

        return self._paragraph


class ReportItemDocxMadrid(ReportItemDocx):
    """Торговая марка (Мадрид)."""
    _paragraph: Paragraph
    document: Document
    obj_type_id: int

    def _write_111(self) -> None:
        """Записывает в документ информацию об ИНИД (111) Номер міжнародної реєстрації."""
        inid = self._get_inid(self.obj_type_id, '111', self.application_data['search_data']['obj_state'])
        registration_number = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('@INTREGN')
        if inid and inid.visible and registration_number:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(registration_number).bold = True
            self._paragraph.add_run('\r')

    def _write_540(self) -> None:
        """Записывает в документ информацию об ИНИД (540) Зображення торговельної марки."""
        inid = self._get_inid(self.obj_type_id, '540', self.application_data['search_data']['obj_state'])
        if inid and inid.visible:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            self._paragraph.add_run('\r')

            # Зображення
            try:
                path = Path(self.application_data['Document']['filesPath'].replace('\\', '/'))
            except KeyError:
                pass
            else:
                parts_len = len(path.parts)
                # Путь к изображению на диске
                mark_image_filepath = Path(settings.MEDIA_ROOT) / path.parts[parts_len - 4].upper()
                mark_image_filepath /= mark_image_filepath / path.parts[parts_len - 3]
                mark_image_filepath /= mark_image_filepath / path.parts[parts_len - 2]
                mark_image_filepath /= path.parts[parts_len - 1]
                registration_number = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('@INTREGN')
                mark_image_filepath /= f"{registration_number}.jpg"
                try:
                    run = self._paragraph.add_run()
                    run.add_picture(str(mark_image_filepath), width=Inches(2.5))
                    self._paragraph.add_run('\r')
                except (UnrecognizedImageError, ZeroDivisionError):
                    pass

    def _write_151(self) -> None:
        """Записывает в документ информацию об ИНИД (151) Дата міжнародної реєстрації."""
        inid = self._get_inid(self.obj_type_id, '151', self.application_data['search_data']['obj_state'])
        registration_date = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('@INTREGD')
        if inid and inid.visible and registration_date:
            registration_date = datetime.strptime(registration_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(registration_date).bold = True
            self._paragraph.add_run('\r')

    def _write_180(self) -> None:
        """Записывает в документ информацию об
        ИНИД (180) Очікувана дата закінчення строку дії реєстрації/продовження."""
        inid = self._get_inid(self.obj_type_id, '180', self.application_data['search_data']['obj_state'])
        exp_date = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('@INTREGD')
        if inid and inid.visible and exp_date:
            exp_date = datetime.strptime(exp_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(exp_date).bold = True
            self._paragraph.add_run('\r')

    def _write_441(self) -> None:
        """Записывает в документ информацию об
        ИНИД (441) Дата публікації відомостей про міжнародну реєстрацію торговельної марки,
        що надійшла для проведення експертизи."""
        inid = self._get_inid(self.obj_type_id, '441', self.application_data['search_data']['obj_state'])
        code_441 = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('Code_441')
        if inid and inid.visible and code_441:
            bul_number = str(bulletin_get_number_441_code(code_441))
            code_441 = datetime.strptime(code_441, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(code_441).bold = True
            if self.lang_code == 'ua':
                self._paragraph.add_run(', бюл. № ').bold = True
            else:
                self._paragraph.add_run(', bul. № ').bold = True
            self._paragraph.add_run(bul_number).bold = True
            self._paragraph.add_run()
            self._paragraph.add_run('\r')

    def _write_450(self) -> None:
        """Записывает в документ информацию об
        ИНИД (450) Дата публікації відомостей про міжнародну реєстрацію та номер бюлетеню Міжнародного бюро ВОІВ"""
        inid = self._get_inid(self.obj_type_id, '450', self.application_data['search_data']['obj_state'])
        pub_date = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('ENN', {}).get('@PUBDATE')
        if inid and inid.visible and pub_date:
            pub_date = datetime.strptime(pub_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            bul_number = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('ENN', {}).get('@GAZNO')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(pub_date).bold = True
            if self.lang_code == 'ua':
                self._paragraph.add_run(', бюл. № ').bold = True
            else:
                self._paragraph.add_run(', bul. № ').bold = True
            self._paragraph.add_run(bul_number).bold = True
            self._paragraph.add_run()
            self._paragraph.add_run('\r')

    def _write_732(self) -> None:
        """Записывает в документ информацию об
        ИНИД (732) Ім'я та адреса володільця реєстрації."""
        inid = self._get_inid(self.obj_type_id, '732', self.application_data['search_data']['obj_state'])
        holder = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('HOLGR')
        if inid and inid.visible and holder:
            holder_name = holder.get('NAME', {}).get('NAMEL')
            holder_address = '\n'.join([x for x in holder.get('ADDRESS', {}).get('ADDRL', [])])
            holder_country = holder.get('ADDRESS', {}).get('COUNTRY')

            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")

            self._paragraph.add_run(holder_name).bold = True
            self._paragraph.add_run('\r')

            self._paragraph.add_run(holder_address)
            self._paragraph.add_run('\r')

            self._paragraph.add_run(f"({holder_country})")
            self._paragraph.add_run('\r')

    def _write_740(self) -> None:
        """Записывает в документ информацию об
        ИНИД (740) Ім'я та адреса представника."""
        inid = self._get_inid(self.obj_type_id, '740', self.application_data['search_data']['obj_state'])
        representative = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('REPGR')
        if inid and inid.visible and representative:
            representative_name = representative.get('NAME', {}).get('NAMEL')
            representative_address = '\n'.join([x for x in representative.get('ADDRESS', {}).get('ADDRL', [])])
            representative_country = representative.get('ADDRESS', {}).get('COUNTRY')

            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")

            self._paragraph.add_run(representative_name).bold = True
            self._paragraph.add_run('\r')

            self._paragraph.add_run(representative_address)
            self._paragraph.add_run('\r')

            self._paragraph.add_run(f"({representative_country})")
            self._paragraph.add_run('\r')

    def _write_821(self) -> None:
        """Записывает в документ информацию об
        ИНИД (821) Базова заявка."""
        inid = self._get_inid(self.obj_type_id, '821', self.application_data['search_data']['obj_state'])
        base_app = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('BASGR', {}).get('BASAPPGR')
        if inid and inid.visible and base_app:
            app_date = datetime.strptime(base_app['BASAPPD'], '%Y-%m-%d').strftime('%d.%m.%Y')

            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(f"{app_date}, ").bold = True
            self._paragraph.add_run(base_app['BASAPPN']).bold = True
            self._paragraph.add_run('\r')

    def _write_891(self) -> None:
        """Записывает в документ информацию об
        ИНИД (891) Дата територіального поширення міжнародної реєстрації."""
        inid = self._get_inid(self.obj_type_id, '891', self.application_data['search_data']['obj_state'])
        reg_date = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('@REGEDAT')
        if inid and inid.visible and reg_date:
            reg_date = datetime.strptime(reg_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(f"{reg_date}, ").bold = True
            self._paragraph.add_run('\r')

    def _write_511(self) -> None:
        """Записывает в документ информацию об
        ИНИД (511) Індекс (індекси) МКТП та перелік товарів і послуг."""
        inid = self._get_inid(self.obj_type_id, '511', self.application_data['search_data']['obj_state'])
        gsgr = self.application_data['MadridTradeMark']['TradeMarkDetails'].get('BASICGS', {}).get('GSGR')
        if inid and inid.visible and gsgr:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for cl in gsgr:
                self._paragraph.add_run(f"\nКл. {cl.get('@NICCLAI')}: ").bold = True
                self._paragraph.add_run(cl.get('GSTERMEN'))
            self._paragraph.add_run('\r')

    def write(self, document: Document) -> Paragraph:
        """Записывает информацию о ТМ в абзац и возвращает его."""
        self.document = document
        self._paragraph = self.document.add_paragraph('')

        self._write_540()
        self._write_111()
        self._write_151()
        self._write_180()
        self._write_441()
        self._write_450()
        self._write_732()
        self._write_740()
        self._write_821()
        self._write_891()
        self._write_511()

        return self._paragraph


class ReportItemDocxMadrid9(ReportItemDocxMadrid):
    obj_type_id = 9


class ReportItemDocxMadrid14(ReportItemDocxMadrid):
    obj_type_id = 14


class ReportItemDocxID(ReportItemDocx):
    """Торговая марка."""
    _paragraph: Paragraph
    document: Document
    obj_type_id = 6

    def _write_21(self) -> None:
        """Записывает в документ данные об ИНИД (21) Номер заявки."""
        inid = self._get_inid(self.obj_type_id, '21', self.application_data['search_data']['obj_state'])
        app_number = self.application_data['search_data']['app_number']
        if inid and inid.visible and app_number:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(app_number).bold = True
            self._paragraph.add_run('\r')

    def _write_51(self) -> None:
        """Записывает в документ данные об
        ИНИД (51) Індекс(и) Міжнародної класифікації промислових зразків (МКПЗ)."""
        inid = self._get_inid(self.obj_type_id, '51', self.application_data['search_data']['obj_state'])
        classes = self.application_data['Design']['DesignDetails'].get('IndicationProductDetails')
        if inid and inid.visible and classes:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(', '.join([x['Class'] for x in classes])).bold = True
            self._paragraph.add_run('\r')

    def _write_11(self) -> None:
        """Записывает в документ данные об
        ИНИД (11) Номер реєстрації (патенту)."""
        inid = self._get_inid(self.obj_type_id, '11', self.application_data['search_data']['obj_state'])
        reg_number = self.application_data['Design']['DesignDetails'].get('RegistrationNumber')
        if inid and inid.visible and reg_number:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(reg_number).bold = True
            self._paragraph.add_run('\r')

    def _write_22(self) -> None:
        """Записывает в документ данные об
        ИНИД (22) Дата подання заявки."""
        inid = self._get_inid(self.obj_type_id, '22', self.application_data['search_data']['obj_state'])
        app_date = self.application_data['Design']['DesignDetails'].get('DesignApplicationDate')
        if inid and inid.visible and app_date:
            app_date = datetime.strptime(app_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(app_date).bold = True
            self._paragraph.add_run('\r')

    def _write_23(self) -> None:
        """Записывает в документ данные об
        ИНИД (23) Дата виставкового пріоритету."""
        inid = self._get_inid(self.obj_type_id, '23', self.application_data['search_data']['obj_state'])
        priority = self.application_data['Design']['DesignDetails'].get('ExhibitionPriorityDetails')
        if inid and inid.visible and priority and len(priority) > 0:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\r")

            for item in priority:
                exhibition_date = datetime.strptime(item['ExhibitionDate'], '%Y-%m-%d').strftime('%d.%m.%Y')
                self._paragraph.add_run(exhibition_date).bold = True
                self._paragraph.add_run('; ')
                self._paragraph.add_run(item['ExhibitionCountryCode'])
                self._paragraph.add_run('\r')

            self._paragraph.add_run('\r')

    def _write_24(self) -> None:
        """Записывает в документ данные об
        ИНИД (24) Дата, з якої є чинними права на промисловий зразок."""
        inid = self._get_inid(self.obj_type_id, '24', self.application_data['search_data']['obj_state'])
        reg_date = self.application_data['Design']['DesignDetails'].get('RecordEffectiveDate')
        if inid and inid.visible and reg_date:
            reg_date = datetime.strptime(reg_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(reg_date).bold = True
            self._paragraph.add_run('\r')

    def _write_28(self) -> None:
        """Записывает в документ данные об
        ИНИД (28) Кількість заявлених варіантів."""
        inid = self._get_inid(self.obj_type_id, '28', self.application_data['search_data']['obj_state'])
        total_specimen = self.application_data['Design']['DesignDetails'].get('TotalSpecimen')
        if inid and inid.visible and total_specimen:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(str(total_specimen)).bold = True
            self._paragraph.add_run('\r')

    def _write_31(self) -> None:
        """Записывает в документ данные об
        ИНИД (31) Номер попередньої заявки відповідно до Паризької конвенції."""
        inid = self._get_inid(self.obj_type_id, '31', self.application_data['search_data']['obj_state'])
        priority = self.application_data['Design']['DesignDetails'].get('PriorityDetails', {}).get('Priority')
        if inid and inid.visible and priority and len(priority) > 0:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(priority[0]['PriorityNumber']).bold = True
            self._paragraph.add_run('\r')

    def _write_32(self) -> None:
        """Записывает в документ данные об
        ИНИД (32) Дата подання попередньої заявки відповідно до Паризької конвенції."""
        inid = self._get_inid(self.obj_type_id, '32', self.application_data['search_data']['obj_state'])
        priority = self.application_data['Design']['DesignDetails'].get('PriorityDetails', {}).get('Priority')
        if inid and inid.visible and priority and len(priority) > 0:
            priority_date = datetime.strptime(priority[0]['PriorityDate'], '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(priority_date).bold = True
            self._paragraph.add_run('\r')

    def _write_33(self) -> None:
        """Записывает в документ данные об
        ИНИД (33) Двобуквений код держави-учасниці Паризької конвенції, до якої подано попередню заявку."""
        inid = self._get_inid(self.obj_type_id, '33', self.application_data['search_data']['obj_state'])
        priority = self.application_data['Design']['DesignDetails'].get('PriorityDetails', {}).get('Priority')
        if inid and inid.visible and priority and len(priority) > 0:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(priority[0]['PriorityCountryCode']).bold = True
            self._paragraph.add_run('\r')

    def _write_45(self) -> None:
        """Записывает в документ данные об ИНИД (45) Дата публікації відомостей про видачу патенту та номер бюлетеня."""
        inid = self._get_inid(self.obj_type_id, '45', self.application_data['search_data']['obj_state'])
        publication_details = self.application_data['Design']['DesignDetails'].get('RecordPublicationDetails')
        if inid and inid.visible and publication_details and len(publication_details) > 0:
            bul_number = publication_details[0]['PublicationIdentifier']
            publication_date = datetime.strptime(
                publication_details[0]['PublicationDate'],
                '%Y-%m-%d'
            ).strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(publication_date).bold = True
            if self.lang_code == 'ua':
                self._paragraph.add_run(', бюл. № ').bold = True
            else:
                self._paragraph.add_run(', bul. № ').bold = True
            self._paragraph.add_run(bul_number).bold = True
            self._paragraph.add_run()
            self._paragraph.add_run('\r')

    def _write_54(self) -> None:
        """Записывает в документ данные об
        ИНИД (54) Назва промислового зразка."""
        inid = self._get_inid(self.obj_type_id, '54', self.application_data['search_data']['obj_state'])
        title = self.application_data['Design']['DesignDetails'].get('DesignTitle')
        if inid and inid.visible and title:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(title).bold = True
            self._paragraph.add_run('\r')

    def _write_55(self) -> None:
        """Записывает в документ данные об
        ИНИД (55) Назва промислового зразка."""
        inid = self._get_inid(self.obj_type_id, '55', self.application_data['search_data']['obj_state'])
        specimen_details = self.application_data['Design']['DesignDetails'].get('DesignSpecimenDetails')
        if inid and inid.visible and specimen_details:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for i, item in enumerate(specimen_details):
                if 'Colors' in item:
                    self._paragraph.add_run('\n')
                    self._paragraph.add_run(f"{i + 1}-й варіант - ")
                    if isinstance(item['Colors'], list):
                        for color in item['Colors']:
                            self._paragraph.add_run(f"{color['Color']}; ").bold = True
                    else:
                        self._paragraph.add_run(f"{item['Colors']['Color']}").bold = True
            self._paragraph.add_run("\r")

            for specimen in specimen_details:
                for image in specimen['DesignSpecimen']:
                    specimen_index = str(image.get('SpecimenIndex', image.get('SpecimenIdentifier', '')))
                    self._paragraph.add_run(specimen_index)
                    self._paragraph.add_run(':\r')
                    path = Path(self.application_data['Document']['filesPath'].replace('\\', '/'))
                    parts_len = len(path.parts)
                    # Путь к изображению на диске
                    mark_image_filepath = Path(settings.MEDIA_ROOT) / path.parts[parts_len - 3]
                    mark_image_filepath /= mark_image_filepath / path.parts[parts_len - 2]
                    mark_image_filepath /= path.parts[parts_len - 1]
                    mark_image_filepath /= image['SpecimenFilename']
                    if mark_image_filepath.exists():
                        try:
                            run = self._paragraph.add_run()
                            run.add_picture(str(mark_image_filepath), width=Inches(2.5))
                            self._paragraph.add_run('\r')
                        except (UnrecognizedImageError, ZeroDivisionError):
                            pass

    def _write_71(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (71) Ім'я (імена) та адреса (адреси) заявника (заявників)
        """
        inid = self._get_inid(self.obj_type_id, '71', self.application_data['search_data']['obj_state'])
        applicants = self.application_data['Design']['DesignDetails'].get(
            'ApplicantDetails', {}
        ).get(
            'Applicant'
        )
        if inid and inid.visible and applicants:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for applicant in applicants:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        applicant['ApplicantAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        applicant['ApplicantAddressBook']['FormattedNameAddress']['Address']['FreeFormatAddress'][
                            'FreeFormatAddressLine']
                    )
                except KeyError:
                    pass

            self._paragraph.add_run('\r')

    def _write_72(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (72) Ім'я (імена) автора (авторів), якщо воно відоме
        """
        inid = self._get_inid(self.obj_type_id, '72', self.application_data['search_data']['obj_state'])
        designers = self.application_data['Design']['DesignDetails'].get(
            'DesignerDetails', {}
        ).get(
            'Designer'
        )
        if inid and inid.visible and designers:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for designer in designers:
                if designer.get('Publicated'):
                    self._paragraph.add_run('\r')
                    try:
                        self._paragraph.add_run(
                            designer['DesignerAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                                'FreeFormatNameDetails']['FreeFormatNameLine']
                        ).bold = True
                    except KeyError:
                        pass

            self._paragraph.add_run('\r')

    def _write_73(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (73) Ім’я або повне найменування власника(ів) патенту, його адреса та двобуквений код держави
        """
        inid = self._get_inid(self.obj_type_id, '73', self.application_data['search_data']['obj_state'])
        holders = self.application_data['Design']['DesignDetails'].get(
            'HolderDetails', {}
        ).get(
            'Holder'
        )
        if inid and inid.visible and holders:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for holder in holders:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        holder['HolderAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        holder['HolderAddressBook']['FormattedNameAddress']['Address']['FreeFormatAddress'][
                            'FreeFormatAddressLine']
                    )
                except KeyError:
                    pass

                self._paragraph.add_run('\r')
                try:
                    country = holder['HolderAddressBook']['FormattedNameAddress']['Address'][
                        'AddressCountryCode']
                    self._paragraph.add_run(
                        f"({country})"
                    )
                except KeyError:
                    pass

            self._paragraph.add_run('\r')

    def _write_74(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (74) Ім'я (імена) та адреса (адреси) представника (представників)
        """
        inid = self._get_inid(self.obj_type_id, '74', self.application_data['search_data']['obj_state'])
        representatives = self.application_data['Design']['DesignDetails'].get(
            'RepresentativeDetails', {}
        ).get(
            'Representative'
        )
        if inid and inid.visible and representatives:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            for representative in representatives:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        representative['RepresentativeAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        representative['RepresentativeAddressBook']['FormattedNameAddress']['Address']['FreeFormatAddress'][
                            'FreeFormatAddressLine']
                    )
                except KeyError:
                    pass

            self._paragraph.add_run('\r')

    def _write_98(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (98) Адреса для листування
        """
        inid = self._get_inid(self.obj_type_id, '98', self.application_data['search_data']['obj_state'])
        correspondence = self.application_data['Design']['DesignDetails'].get(
            'CorrespondenceAddress'
        )
        if inid and inid.visible and correspondence:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:")
            self._paragraph.add_run('\r')
            try:
                self._paragraph.add_run(
                    correspondence['CorrespondenceAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                        'FreeFormatNameDetails']['FreeFormatNameLine']
                ).bold = True
            except KeyError:
                pass

            self._paragraph.add_run('\r')
            try:
                self._paragraph.add_run(
                    correspondence['CorrespondenceAddressBook']['FormattedNameAddress']['Address']['FreeFormatAddress'][
                        'FreeFormatAddressLine']
                )
            except KeyError:
                pass

            self._paragraph.add_run('\r')
            try:
                country = correspondence['CorrespondenceAddressBook']['FormattedNameAddress']['Address'][
                    'AddressCountryCode']
                self._paragraph.add_run(
                    f"({country})"
                )
            except KeyError:
                pass

            self._paragraph.add_run('\r')

    def write(self, document: Document) -> Paragraph:
        """Записывает информацию о ТМ в абзац и возвращает его."""
        self.document = document
        self._paragraph = self.document.add_paragraph('')

        self._write_51()
        self._write_11()
        self._write_21()
        self._write_22()
        self._write_23()
        self._write_24()
        self._write_31()
        self._write_32()
        self._write_33()
        self._write_45()
        self._write_54()
        self._write_55()
        self._write_71()
        self._write_72()
        self._write_73()
        self._write_74()
        self._write_98()

        self._write_28()

        return self._paragraph


class ReportItemDocxGeo(ReportItemDocx):
    """Торговая марка."""
    _paragraph: Paragraph
    document: Document
    obj_type_id = 5

    def _write_111(self) -> None:
        """Записывает в документ данные об
        ИНИД (111) Номер реєстрації."""
        inid = self._get_inid(self.obj_type_id, '111', self.application_data['search_data']['obj_state'])
        reg_number = self.application_data['Geo']['GeoDetails'].get('RegistrationNumber')
        if inid and inid.visible and reg_number:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(reg_number).bold = True
            self._paragraph.add_run('\r')

    def _write_151(self) -> None:
        """Записывает в документ данные об
        ИНИД (151) Номер реєстрації."""
        inid = self._get_inid(self.obj_type_id, '151', self.application_data['search_data']['obj_state'])
        reg_date = self.application_data['Geo']['GeoDetails'].get('RegistrationDate')
        if inid and inid.visible and reg_date:
            reg_date = datetime.strptime(reg_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(reg_date).bold = True
            self._paragraph.add_run('\r')

    def _write_190(self) -> None:
        """Записывает в документ данные об
        ИНИД (190) Держава реєстрації КЗПТ."""
        inid = self._get_inid(self.obj_type_id, '190', self.application_data['search_data']['obj_state'])
        country = self.application_data['Geo']['GeoDetails'].get('RegistrationOriginCountry', [])
        if inid and inid.visible and len(country) > 0:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(country[0]).bold = True
            self._paragraph.add_run('\r')

    def _write_210(self) -> None:
        """Записывает в документ данные об
        ИНИД (210) Номер заявки."""
        inid = self._get_inid(self.obj_type_id, '210', self.application_data['search_data']['obj_state'])
        app_number = self.application_data['Geo']['GeoDetails'].get('ApplicationNumber', [])
        if inid and inid.visible and app_number:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(app_number).bold = True
            self._paragraph.add_run('\r')

    def _write_220(self) -> None:
        """Записывает в документ данные об
        ИНИД (220) Дата подання заявки."""
        inid = self._get_inid(self.obj_type_id, '220', self.application_data['search_data']['obj_state'])
        app_date = self.application_data['Geo']['GeoDetails'].get('ApplicationDate')
        if inid and inid.visible and app_date:
            app_date = datetime.strptime(app_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(app_date).bold = True
            self._paragraph.add_run('\r')

    def _write_539i(self) -> None:
        """Записывает в документ данные об
        ИНИД (539.I) Назва КЗПТ."""
        inid = self._get_inid(self.obj_type_id, '539.I', self.application_data['search_data']['obj_state'])
        product_name = self.application_data['Geo']['GeoDetails'].get('ProductName')
        if inid and inid.visible and product_name:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(product_name).bold = True
            self._paragraph.add_run('\r')

    def _write_540(self) -> None:
        """Записывает в документ данные об
        ИНИД (540) Назва товару."""
        inid = self._get_inid(self.obj_type_id, '540', self.application_data['search_data']['obj_state'])
        indication = self.application_data['Geo']['GeoDetails'].get('Indication')
        if inid and inid.visible and indication:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(indication).bold = True
            self._paragraph.add_run('\r')

    def _write_732(self) -> None:
        """Записывает в документ данные об
        ИНИД (732) Ім'я та адреса володільця реєстрації."""
        inid = self._get_inid(self.obj_type_id, '732', self.application_data['search_data']['obj_state'])
        holders = self.application_data['Geo']['GeoDetails'].get('HolderDetails', {}).get('Holder')
        if inid and inid.visible and holders:
            self._paragraph.add_run(f"{inid.title}:\n")
            for holder in holders:
                try:
                    self._paragraph.add_run(
                        holder['HolderAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        holder['HolderAddressBook']['FormattedNameAddress']['Address']['FreeFormatAddress'][
                            'FreeFormatAddressLine']
                    )
                except KeyError:
                    pass

                try:
                    country = holder['HolderAddressBook']['FormattedNameAddress']['Address']['AddressCountryCode']
                    self._paragraph.add_run(
                        f" ({country})"
                    )
                except KeyError:
                    pass
            self._paragraph.add_run('\r')

    def _write_750(self) -> None:
        """
        Записывает в документ
        данные об ИНИД (750) Адреса для листування
        """
        inid = self._get_inid(self.obj_type_id, '750', self.application_data['search_data']['obj_state'])
        correspondence = self.application_data['Geo']['GeoDetails'].get(
            'CorrespondenceAddress', {}
        ).get(
            'CorrespondenceAddressBook', {}
        ).get(
            'FormattedNameAddress'
        )
        if inid and inid.visible and correspondence:
            self._paragraph.add_run(f"{inid.title}:")
            self._paragraph.add_run('\r')
            try:
                self._paragraph.add_run(
                    correspondence['Name']['FreeFormatName']['FreeFormatNameDetails']['FreeFormatNameLine']
                ).bold = True
            except KeyError:
                pass

            self._paragraph.add_run('\r')
            try:
                self._paragraph.add_run(
                    correspondence['Address']['FreeFormatAddress']['FreeFormatAddressLine']
                )
            except KeyError:
                pass

            try:
                country = correspondence['Address']['AddressCountryCode']
                self._paragraph.add_run(
                    f" ({country})"
                )
            except KeyError:
                pass
        self._paragraph.add_run('\r')

    def write(self, document: Document) -> Paragraph:
        """Записывает информацию о геогр в абзац и возвращает его."""
        self.document = document
        self._paragraph = self.document.add_paragraph('')

        self._write_111()
        self._write_151()
        self._write_190()
        self._write_210()
        self._write_220()
        self._write_539i()
        self._write_540()
        self._write_732()
        self._write_750()

        return self._paragraph


class ReportItemInvUMLD(ReportItemDocx):
    """Изобретения, полезные модели."""
    _paragraph: Paragraph
    document: Document
    obj_type_id: int
    body: dict

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        try:
            if self.application_data['search_data']['obj_state'] == 1:
                self.body = self.application_data['Claim']
            else:
                self.body = self.application_data['Patent']
        except KeyError:
            self.body = {}

    def _write_11(self):
        """Записывает в документ данные об
        ИНИД (11) Номер патенту"""
        inid = self._get_inid(self.obj_type_id, '11', self.application_data['search_data']['obj_state'])
        reg_number = self.body.get('I_11')
        if inid and inid.visible and reg_number:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(str(reg_number)).bold = True
            self._paragraph.add_run('\r')

    def _write_21(self):
        """Записывает в документ данные об
        ИНИД (21) Номер заявки"""
        inid = self._get_inid(self.obj_type_id, '21', self.application_data['search_data']['obj_state'])
        app_number = self.application_data['search_data']['app_number']
        if inid and inid.visible and app_number:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(app_number).bold = True
            self._paragraph.add_run('\r')

    def _write_22(self):
        """Записывает в документ данные об
        ИНИД (22) Дата подання заявки"""
        inid = self._get_inid(self.obj_type_id, '22', self.application_data['search_data']['obj_state'])
        app_date = self.body.get('I_22')
        if inid and inid.visible and app_date:
            app_date = datetime.strptime(app_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(app_date).bold = True
            self._paragraph.add_run('\r')

    def _write_24(self):
        """Записывает в документ данные об
        ИНИД (24) Дата, з якої є чинними права"""
        inid = self._get_inid(self.obj_type_id, '24', self.application_data['search_data']['obj_state'])
        reg_date = self.body.get('I_24')
        if inid and inid.visible and reg_date:
            reg_date = datetime.strptime(reg_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(reg_date).bold = True
            self._paragraph.add_run('\r')

    def _write_31(self):
        """Записывает в документ данные об
        ИНИД (31) Номер попередньої заявки"""
        inid = self._get_inid(self.obj_type_id, '31', self.application_data['search_data']['obj_state'])
        i_30 = self.body.get('I_30')
        if inid and inid.visible and i_30:
            for item in i_30:
                i_31 = item.get('I_31')
                if i_31:
                    self._paragraph.add_run(f"({inid.code})").bold = True
                    self._paragraph.add_run(f"\t{inid.title}: ")
                    self._paragraph.add_run(i_31).bold = True
                    self._paragraph.add_run('\r')

    def _write_32(self):
        """Записывает в документ данные об
        ИНИД (32) Дата подання попередньої заявки"""
        inid = self._get_inid(self.obj_type_id, '32', self.application_data['search_data']['obj_state'])
        i_30 = self.body.get('I_30')
        if inid and inid.visible and i_30:
            for item in i_30:
                i_32 = item.get('I_32')
                if i_32:
                    i_32 = datetime.strptime(i_32, '%Y-%m-%d').strftime('%d.%m.%Y')
                    self._paragraph.add_run(f"({inid.code})").bold = True
                    self._paragraph.add_run(f"\t{inid.title}: ")
                    self._paragraph.add_run(i_32).bold = True
                    self._paragraph.add_run('\r')

    def _write_33(self):
        """Записывает в документ данные об
        ИНИД (33) Двобуквений код держави"""
        inid = self._get_inid(self.obj_type_id, '33', self.application_data['search_data']['obj_state'])
        i_30 = self.body.get('I_30')
        if inid and inid.visible and i_30:
            for item in i_30:
                i_33 = item.get('I_33')
                if i_33:
                    self._paragraph.add_run(f"({inid.code})").bold = True
                    self._paragraph.add_run(f"\t{inid.title}: ")
                    self._paragraph.add_run(i_33).bold = True
                    self._paragraph.add_run('\r')

    def _write_41(self):
        """Записывает в документ данные об
        ИНИД (41) Дата публікації заявки"""
        inid = self._get_inid(self.obj_type_id, '41', self.application_data['search_data']['obj_state'])
        i_43 = self.body.get('I_43.D')
        if inid and inid.visible and i_43:
            for item in i_43:
                value = datetime.strptime(item, '%Y-%m-%d').strftime('%d.%m.%Y')
                self._paragraph.add_run(f"({inid.code})").bold = True
                self._paragraph.add_run(f"\t{inid.title}: ")
                self._paragraph.add_run(value).bold = True
                self._paragraph.add_run('\r')

    def _write_46(self):
        """Записывает в документ данные об
        ИНИД (46) Дата публікації відомостей про видачу патенту, номер бюлетня"""
        inid = self._get_inid(self.obj_type_id, '46', self.application_data['search_data']['obj_state'])
        i_45 = self.body.get('I_45.D')
        if inid and inid.visible and i_45:
            for item in i_45:
                value = datetime.strptime(item, '%Y-%m-%d').strftime('%d.%m.%Y')
                self._paragraph.add_run(f"({inid.code})").bold = True
                self._paragraph.add_run(f"\t{inid.title}: ")
                self._paragraph.add_run(value).bold = True
                if self.body.get('I_45_bul_str'):
                    if self.lang_code == 'ua':
                        self._paragraph.add_run(f", бюл. № {self.body.get('I_45_bul_str')}").bold = True
                    else:
                        self._paragraph.add_run(f", bul. № {self.body.get('I_45_bul_str')}").bold = True
                self._paragraph.add_run('\r')

    def _write_51(self):
        """Записывает в документ данные об
        ИНИД (51) Iндекс МПК"""
        inid = self._get_inid(self.obj_type_id, '51', self.application_data['search_data']['obj_state'])
        ipc = self.body.get('IPC')
        if inid and inid.visible and ipc:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")
            for item in ipc:
                self._paragraph.add_run(item).bold = True
                self._paragraph.add_run('\r')

    def _write_54(self):
        """Записывает в документ данные об
        ИНИД (54) Назва винаходу (корисної моделі)"""
        inid = self._get_inid(self.obj_type_id, '54', self.application_data['search_data']['obj_state'])
        i_54 = self.body.get('I_54')
        if inid and inid.visible and i_54:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")
            for item in i_54:
                try:
                    values = list(item.values())
                    keys = list(item.keys())
                    self._paragraph.add_run(values[0]).bold = True
                    if 'U' in keys[0]:
                        self._paragraph.add_run(' [UA]')
                    elif 'E' in keys[0]:
                        self._paragraph.add_run(' [EN]')
                    elif 'R' in keys[0]:
                        self._paragraph.add_run(' [RU]')
                    self._paragraph.add_run('\r')
                except IndexError:
                    pass

    def _write_56(self):
        """Записывает в документ данные об
        ИНИД (56) Аналоги винаходу"""
        inid = self._get_inid(self.obj_type_id, '56', self.application_data['search_data']['obj_state'])
        i_56 = self.body.get('I_56')
        if inid and inid.visible and i_56:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")
            for item in i_56:
                value = item.get('I_56.A')
                if value:
                    self._paragraph.add_run(value).bold = True
                    self._paragraph.add_run('\r')

    def _write_85(self):
        """Записывает в документ данные об
        ИНИД (85) Дата входження до національної фази"""
        inid = self._get_inid(self.obj_type_id, '85', self.application_data['search_data']['obj_state'])
        i_85 = self.body.get('I_85')
        if inid and inid.visible and i_85:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            i_85 = datetime.strptime(i_85, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(i_85).bold = True
            self._paragraph.add_run('\r')

    def _write_86(self):
        """Записывает в документ данные об
        ИНИД (86) Номер та дата подання заявки РСТ"""
        inid = self._get_inid(self.obj_type_id, '86', self.application_data['search_data']['obj_state'])
        i_86 = self.body.get('I_86')
        if inid and inid.visible and i_86:
            for item in i_86:
                self._paragraph.add_run(f"({inid.code})").bold = True
                self._paragraph.add_run(f"\t{inid.title}:\n")
                i_86_number = item.get('I_86.N')
                if i_86_number:
                    self._paragraph.add_run(i_86_number).bold = True
                i_86_date = item.get('I_86.D')
                if i_86_date:
                    i_86_date = datetime.strptime(i_86_date, '%Y-%m-%d').strftime('%d.%m.%Y')
                    self._paragraph.add_run(f", {i_86_date}").bold = True
                self._paragraph.add_run('\r')

    def _write_71(self):
        """Записывает в документ данные об
        ИНИД (71) Заявник"""
        inid = self._get_inid(self.obj_type_id, '71', self.application_data['search_data']['obj_state'])
        i_71 = self.body.get('I_71')
        if inid and inid.visible and i_71:
            i_71 = list_of_dicts_unique(i_71)
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")
            for item in i_71:
                person_name = get_person_name(item)
                person_country = get_person_country(item)
                self._paragraph.add_run(person_name).bold = True
                self._paragraph.add_run(f" ({person_country})")
                self._paragraph.add_run('\r')

    def _write_72(self):
        """Записывает в документ данные об
        ИНИД (72) Винахідник"""
        inid = self._get_inid(self.obj_type_id, '72', self.application_data['search_data']['obj_state'])
        i_72 = self.body.get('I_72')
        if inid and inid.visible and i_72:
            i_72 = list_of_dicts_unique(i_72)
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")
            for item in i_72:
                person_name = get_person_name(item)
                person_country = get_person_country(item)
                self._paragraph.add_run(person_name).bold = True
                self._paragraph.add_run(f" ({person_country})")
                self._paragraph.add_run('\r')

    def _write_73(self):
        """Записывает в документ данные об
        ИНИД (73) Власник"""
        inid = self._get_inid(self.obj_type_id, '73', self.application_data['search_data']['obj_state'])
        i_73 = self.body.get('I_73')
        if inid and inid.visible and i_73:
            i_73 = list_of_dicts_unique(i_73)
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}:\n")
            for item in i_73:
                person_name = get_person_name(item)
                person_country = get_person_country(item)
                self._paragraph.add_run(person_name).bold = True
                self._paragraph.add_run(f" ({person_country})")
                self._paragraph.add_run('\r')

    def _write_74(self):
        """Записывает в документ данные об
        ИНИД (74) Представник"""
        inid = self._get_inid(self.obj_type_id, '74', self.application_data['search_data']['obj_state'])
        i_74 = self.body.get('I_74')
        if inid and inid.visible and i_74:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(i_74).bold = True
            self._paragraph.add_run('\r')

    def _write_98(self):
        """Записывает в документ данные об
        ИНИД (98) Адреса для листування"""
        inid = self._get_inid(self.obj_type_id, '98', self.application_data['search_data']['obj_state'])
        i_98 = self.body.get('I_98')
        if inid and inid.visible and i_98:
            self._paragraph.add_run(f"({inid.code})").bold = True
            self._paragraph.add_run(f"\t{inid.title}: ")
            self._paragraph.add_run(i_98).bold = True
            self._paragraph.add_run('\r')

    def _write_62(self):
        """Записывает в документ данные об
        ИНИД (62) Номер та дата більш ранньої заявки, з якої було виділено даний патентний документ"""
        inid = self._get_inid(self.obj_type_id, '62', self.application_data['search_data']['obj_state'])
        i_62 = self.body.get('i_62')
        if inid and inid.visible and i_62:
            for item in i_62:
                self._paragraph.add_run(f"({inid.code})").bold = True
                self._paragraph.add_run(f"\t{inid.title}:\n")
                i_62_number = item.get('I_62.N')
                self._paragraph.add_run(i_62_number).bold = True
                i_62_date = item.get('I_62.D')
                if i_62_date:
                    i_62_date = datetime.strptime(i_62_date, '%Y-%m-%d').strftime('%d.%m.%Y')
                    self._paragraph.add_run(f", {i_62_date}").bold = True
                self._paragraph.add_run('\r')

    def write(self, document: Document) -> Paragraph:
        self.document = document
        self._paragraph = self.document.add_paragraph('')

        self._write_11()
        self._write_21()
        self._write_22()
        self._write_24()
        self._write_31()
        self._write_32()
        self._write_33()
        self._write_41()
        self._write_46()
        self._write_51()
        self._write_54()
        self._write_56()
        self._write_62()
        self._write_85()
        self._write_86()
        self._write_71()
        self._write_72()
        self._write_73()
        self._write_74()
        self._write_98()

        return self._paragraph


class ReportItemInv(ReportItemInvUMLD):
    obj_type_id = 1


class ReportItemUM(ReportItemInvUMLD):
    obj_type_id = 2


class ReportItemLD(ReportItemInvUMLD):
    obj_type_id = 3


class ReportItemCopyright(ReportItemDocx):
    """Авторське право на твір"""
    _paragraph: Paragraph
    document: Document
    obj_type_id = 10

    def _write_11(self) -> None:
        """Записывает в документ данные об
        ИНИД (11) Номер свідоцтва про реєстрацію авторського права на твір."""
        inid = self._get_inid(self.obj_type_id, '11', self.application_data['search_data']['obj_state'])
        reg_number = self.application_data['Certificate']['CopyrightDetails'].get('RegistrationNumber')
        if inid and inid.visible and reg_number:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(reg_number).bold = True
            self._paragraph.add_run('\r')

    def _write_15(self) -> None:
        """Записывает в документ данные об
        ИНИД (15) Дата реєстрації авторського права."""
        inid = self._get_inid(self.obj_type_id, '15', self.application_data['search_data']['obj_state'])
        reg_date = self.application_data['Certificate']['CopyrightDetails'].get('RegistrationDate')
        if inid and inid.visible and reg_date:
            reg_date = datetime.strptime(reg_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(reg_date).bold = True
            self._paragraph.add_run('\r')

    def _write_29(self) -> None:
        """Записывает в документ данные об
        ИНИД (29) Об'єкт авторського права, до якого належить твір."""
        inid = self._get_inid(self.obj_type_id, '29', self.application_data['search_data']['obj_state'])
        obj_kind = self.application_data['Certificate']['CopyrightDetails'].get(
            'CopyrightObjectKindDetails', {}
        ).get('CopyrightObjectKind')
        if inid and inid.visible and obj_kind and len(obj_kind) > 0:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(obj_kind[0]['CopyrightObjectKindName']).bold = True
            self._paragraph.add_run('\r')

    def _write_45d(self) -> None:
        """Записывает в документ данные об
        ИНИД (45) Дата публікації"""
        inid = self._get_inid(self.obj_type_id, '45.D', self.application_data['search_data']['obj_state'])
        pub_date = self.application_data['Certificate']['CopyrightDetails'].get(
            'PublicationDetails', {}
        ).get(
            'Publication', {}
        ).get('PublicationDate')
        if inid and inid.visible and pub_date:
            pub_date = datetime.strptime(pub_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(pub_date).bold = True
            self._paragraph.add_run('\r')

    def _write_45n(self) -> None:
        """Записывает в документ данные об
        ИНИД (45) Номер бюлетеня"""
        inid = self._get_inid(self.obj_type_id, '45.N', self.application_data['search_data']['obj_state'])
        pub_number = self.application_data['Certificate']['CopyrightDetails'].get(
            'PublicationDetails', {}
        ).get(
            'Publication', {}
        ).get('PublicationNumber')
        if inid and inid.visible and pub_number:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(pub_number).bold = True
            self._paragraph.add_run('\r')

    def _write_54(self) -> None:
        """Записывает в документ данные об
        ИНИД (54) Вид, повна та скорочена назва твору (творів)"""
        inid = self._get_inid(self.obj_type_id, '54', self.application_data['search_data']['obj_state'])
        name = self.application_data['Certificate']['CopyrightDetails'].get('Name')
        if inid and inid.visible and name:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(name).bold = True
            self._paragraph.add_run('\r')

    def _write_57(self) -> None:
        """Записывает в документ данные об
        ИНИД (57) Анотація"""
        inid = self._get_inid(self.obj_type_id, '57', self.application_data['search_data']['obj_state'])
        annotation = self.application_data['Certificate']['CopyrightDetails'].get('Annotation')
        if inid and inid.visible and annotation:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(annotation).bold = True
            self._paragraph.add_run('\r')

    def _write_58(self) -> None:
        """Записывает в документ данные об
        ИНИД (58) Вихідні дані для оприлюднених творів"""
        inid = self._get_inid(self.obj_type_id, '58', self.application_data['search_data']['obj_state'])
        promulgation = self.application_data['Certificate']['CopyrightDetails'].get('PromulgationData')
        if inid and inid.visible and promulgation:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(promulgation).bold = True
            self._paragraph.add_run('\r')

    def _write_72(self) -> None:
        """Записывает в документ данные об
        ИНИД (72) Повне ім'я та/або псевдонім автора (авторів), чи позначення «Анонімно»"""
        inid = self._get_inid(self.obj_type_id, '72', self.application_data['search_data']['obj_state'])
        authors = self.application_data['Certificate']['CopyrightDetails'].get(
            'AuthorDetails', {}
        ).get(
            'Author'
        )
        if inid and inid.visible and authors:
            self._paragraph.add_run(f"{inid.title}:")
            for author in authors:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        author['AuthorAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                try:
                    country = author['AuthorAddressBook']['FormattedNameAddress']['Address'][
                        'AddressCountryCode']
                    self._paragraph.add_run(
                        f" ({country})"
                    )
                except KeyError:
                    pass
            self._paragraph.add_run('\r')

    def _write_77(self) -> None:
        """Записывает в документ данные об
        ИНИД (77) Повне ім'я або повне офіційне найменування роботодавця"""
        inid = self._get_inid(self.obj_type_id, '77', self.application_data['search_data']['obj_state'])
        employers = self.application_data['Certificate']['CopyrightDetails'].get(
            'EmployerDetails', {}
        ).get(
            'Employer'
        )
        if inid and inid.visible and employers:
            self._paragraph.add_run(f"{inid.title}:")
            for employer in employers:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        employer['EmployerAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                try:
                    country = employer['EmployerAddressBook']['FormattedNameAddress']['Address'][
                        'AddressCountryCode']
                    self._paragraph.add_run(
                        f" ({country})"
                    )
                except KeyError:
                    pass
            self._paragraph.add_run('\r')

    def write(self, document: Document) -> Paragraph:
        self.document = document
        self._paragraph = self.document.add_paragraph('')

        self._write_11()
        self._write_15()
        self._write_29()
        self._write_45d()
        self._write_45n()
        self._write_54()
        self._write_57()
        self._write_58()
        self._write_72()
        self._write_77()

        return self._paragraph


class ReportItemCopyrightOfficialWork(ReportItemCopyright):
    """Авторське право на службовий твір."""
    obj_type_id = 13


class ReportItemAgreement(ReportItemDocx):
    """Авторське право (договір)"""
    _paragraph: Paragraph
    document: Document
    obj_type_id = 11

    def _write_11(self) -> None:
        """Записывает в документ данные об
        ИНИД (11) Номер реєстрації договору."""
        inid = self._get_inid(self.obj_type_id, '11', self.application_data['search_data']['obj_state'])
        reg_number = self.application_data['Decision']['DecisionDetails'].get('RegistrationNumber')
        if inid and inid.visible and reg_number:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(reg_number).bold = True
            self._paragraph.add_run('\r')

    def _write_15(self) -> None:
        """Записывает в документ данные об
        ИНИД (15) Дата реєстрації договору."""
        inid = self._get_inid(self.obj_type_id, '15', self.application_data['search_data']['obj_state'])
        reg_date = self.application_data['Decision']['DecisionDetails'].get('RegistrationDate')
        if inid and inid.visible and reg_date:
            reg_date = datetime.strptime(reg_date, '%Y-%m-%d').strftime('%d.%m.%Y')
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(reg_date).bold = True
            self._paragraph.add_run('\r')

    def _write_27(self) -> None:
        """Записывает в документ данные об
        ИНИД (27) Вид договору."""
        inid = self._get_inid(self.obj_type_id, '27', self.application_data['search_data']['obj_state'])
        reg_kind = self.application_data['Decision']['DecisionDetails'].get('RegistrationKind')
        if inid and inid.visible and reg_kind:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(reg_kind).bold = True
            self._paragraph.add_run('\r')

    def _write_29(self) -> None:
        """Записывает в документ данные об
        ИНИД (29) Об'єкт авторського права, до якого належить твір."""
        inid = self._get_inid(self.obj_type_id, '29', self.application_data['search_data']['obj_state'])
        obj_kind = self.application_data['Decision']['DecisionDetails'].get(
            'CopyrightObjectKindDetails', {}
        ).get('CopyrightObjectKind')
        if inid and inid.visible and obj_kind and len(obj_kind) > 0:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(obj_kind[0]['CopyrightObjectKindName']).bold = True
            self._paragraph.add_run('\r')

    def _write_54(self) -> None:
        """Записывает в документ данные об
        ИНИД (54) Вид, повна та скорочена назва твору (творів)"""
        inid = self._get_inid(self.obj_type_id, '54', self.application_data['search_data']['obj_state'])
        name = self.application_data['Decision']['DecisionDetails'].get('Name')
        if inid and inid.visible and name:
            self._paragraph.add_run(f"{inid.title}: ")
            self._paragraph.add_run(name).bold = True
            self._paragraph.add_run('\r')

    def _write_72(self) -> None:
        """Записывает в документ данные об
        ИНИД (72) Повне ім'я та/або псевдонім автора (авторів), чи позначення «Анонімно»"""
        inid = self._get_inid(self.obj_type_id, '72', self.application_data['search_data']['obj_state'])
        authors = self.application_data['Decision']['DecisionDetails'].get(
            'AuthorDetails', {}
        ).get(
            'Author'
        )
        if inid and inid.visible and authors:
            self._paragraph.add_run(f"{inid.title}:")
            for author in authors:
                self._paragraph.add_run('\r')
                try:
                    self._paragraph.add_run(
                        author['AuthorAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                except KeyError:
                    pass

                try:
                    country = author['AuthorAddressBook']['FormattedNameAddress']['Address'][
                        'AddressCountryCode']
                    self._paragraph.add_run(
                        f" ({country})"
                    )
                except KeyError:
                    pass
            self._paragraph.add_run('\r')

    def _write_75(self) -> None:
        """Записывает в документ данные об
        ИНИД (75) Повне ім'я фізичної(их) або повне офіційне найменування юридичної(их) особи (осіб), сторін договору"""
        inid = self._get_inid(self.obj_type_id, '75', self.application_data['search_data']['obj_state'])

        licensors = self.application_data['Decision']['DecisionDetails'].get(
            'LicensorDetails', {}
        ).get(
            'Licensor'
        )
        licensees = self.application_data['Decision']['DecisionDetails'].get(
            'LicenseeDetails', {}
        ).get(
            'Licensee'
        )

        if inid and inid.visible and licensors and licensees:
            self._paragraph.add_run(f"{inid.title}:\n")

            for i, licensor in enumerate(licensors):
                try:
                    self._paragraph.add_run(
                        licensor['LicensorAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                    if i < len(licensors) - 1:
                        self._paragraph.add_run('; ')
                except KeyError:
                    pass
            self._paragraph.add_run(' - ')
            for i, licensee in enumerate(licensees):
                try:
                    self._paragraph.add_run(
                        licensee['LicenseeAddressBook']['FormattedNameAddress']['Name']['FreeFormatName'][
                            'FreeFormatNameDetails']['FreeFormatNameLine']
                    ).bold = True
                    if i < len(licensors) - 1:
                        self._paragraph.add_run('; ')
                except KeyError:
                    pass

    def write(self, document: Document) -> Paragraph:
        self.document = document
        self._paragraph = self.document.add_paragraph('')

        self._write_11()
        self._write_15()
        self._write_27()
        self._write_29()
        self._write_54()
        self._write_72()
        self._write_75()

        return self._paragraph


class ReportItemAgreementTransfer(ReportItemAgreement):
    obj_type_id = 12


class ReportWriter(ABC):
    """Интерфейс создателя файла отчёта."""
    items: List[ReportItem]

    def __init__(self, items: List[ReportItem]):
        self.items = items

    @abstractmethod
    def generate(self, file_path: Path) -> Path:
        raise NotImplemented


class ReportWriterDocx(ReportWriter):
    """Создатель файла отчёта в формате docx."""
    items: List[ReportItemDocx]
    _obj_type_titles = [
        {
            'obj_type_id': 1,
            'obj_state': 1,
            'ua': 'Заявка на винахід',
            'en': 'Invention Application',
        },
        {
            'obj_type_id': 1,
            'obj_state': 2,
            'ua': 'Патент на винахід',
            'en': 'Invention Patent ',
        },
        {
            'obj_type_id': 2,
            'obj_state': 1,
            'ua': 'Заявка на корисну модель',
            'en': 'Utility Model Application',
        },
        {
            'obj_type_id': 2,
            'obj_state': 2,
            'ua': 'Патент на корисну модель',
            'en': 'Utility Model Patent',
        },
        {
            'obj_type_id': 3,
            'obj_state': 2,
            'ua': 'Свідоцтво України на топографію (компонування) ІМС',
            'en': 'Certificate of Ukraine on topography (layout) of IC',
        },
        {
            'obj_type_id': 4,
            'obj_state': 1,
            'ua': 'Заявка на торговельну марку',
            'en': 'Trademark Application',
        },
        {
            'obj_type_id': 4,
            'obj_state': 2,
            'ua': 'Свідоцтво України на торговельну марку',
            'en': 'Trademark certificate of Ukraine',
        },
        {
            'obj_type_id': 5,
            'obj_state': 2,
            'ua': 'Свідоцтво на використання географічного зазначення',
            'en': 'Certificate of using of a geographical indications',
        },
        {
            'obj_type_id': 6,
            'obj_state': 1,
            'ua': 'Заявка на промисловий зразок',
            'en': 'Industrial design Application',
        },
        {
            'obj_type_id': 6,
            'obj_state': 2,
            'ua': 'Патент України на промисловий зразок',
            'en': 'Patent of Ukraine for industrial design',
        },
        {
            'obj_type_id': 9,
            'obj_state': 2,
            'ua': 'Міжнародна реєстрація торговельної марки з поширенням на територію України',
            'en': 'International trademark registration with distribution to the territory of Ukraine',
        },
        {
            'obj_type_id': 10,
            'obj_state': 2,
            'ua': 'Авторське право на твір',
            'en': 'Copyright to the work',
        },
        {
            'obj_type_id': 11,
            'obj_state': 2,
            'ua': 'Договір про передачу права на використання твору',
            'en': 'Agreement on the transfer of the right to use the work',
        },
        {
            'obj_type_id': 12,
            'obj_state': 2,
            'ua': 'Договір про передачу (відчуження) майнових прав на твір',
            'en': 'Agreement on the transfer (alienation) of property rights to the work',
        },
        {
            'obj_type_id': 13,
            'obj_state': 2,
            'ua': 'Авторське право на службовий твір',
            'en': 'Copyright for an official work',
        },
        {
            'obj_type_id': 14,
            'obj_state': 2,
            'ua': 'Міжнародна реєстрація торговельної марки, що зареєстрована в Україні',
            'en': 'International trademark registration registered in Ukraine',
        },
    ]

    def _set_font(self, document: Document(), font_name='Times New Roman'):
        style = document.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(12)

        for paragraph in document.paragraphs:
            paragraph.style = document.styles['Normal']

    def _get_obj_type_title(self, obj_type_id: int, obj_state: int, lang_code: str) -> str:
        for item in self._obj_type_titles:
            if item['obj_type_id'] == obj_type_id and item['obj_state'] == obj_state:
                return item[lang_code]
        return ''

    def generate(self, file_path: Path):
        document = Document()
        for i, item in enumerate(self.items):
            p = document.add_paragraph()
            p.add_run(f"{str(i + 1)}. ").bold = True
            obj_type_title = self._get_obj_type_title(
                item.application_data['Document']['idObjType'],
                item.application_data['search_data']['obj_state'],
                item.lang_code,
            )
            if obj_type_title:
                p.add_run(f" {obj_type_title}").bold = True
            item.write(document)
        self._set_font(document)
        document.save(str(file_path))


class ReportWriterCreator(ABC):
    """Интерфейс создателя генератора отчётов."""
    @staticmethod
    @abstractmethod
    def create(applications: List[dict], inid_data: List[InidCode], lang_code: str) -> ReportWriter:
        """
        Создаёт объект генератора отчётов.

        :param applications список заявок
        :param inid_data список кодов инид
        :param lang_code языковой код
        """
        raise NotImplemented


class ReportWriterDocxCreator(ReportWriterCreator):
    """Класс, задачей которого есть создание объекта генератора отчётов в формате .docx."""
    @staticmethod
    def create(applications: List[dict], inid_data: List[InidCode], lang_code: str = 'ua') -> ReportWriterDocx:
        report_item_classes = {
            1: ReportItemInv,
            2: ReportItemUM,
            3: ReportItemLD,
            4: ReportItemDocxTM,
            5: ReportItemDocxGeo,
            6: ReportItemDocxID,
            9: ReportItemDocxMadrid9,
            10: ReportItemCopyright,
            11: ReportItemAgreement,
            12: ReportItemAgreementTransfer,
            13: ReportItemCopyrightOfficialWork,
            14: ReportItemDocxMadrid14,
        }
        report_items = []
        for app in applications:
            report_item = report_item_classes[app['Document']['idObjType']](
                application_data=app,
                ipc_fields=inid_data,
                lang_code=lang_code,
            )
            report_items.append(report_item)
        return ReportWriterDocx(items=report_items)

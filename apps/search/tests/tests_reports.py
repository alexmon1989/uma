from django.test import TestCase
from apps.search.services.reports import ReportItemDocxTM, ReportItemDocxID, ReportItemDocxMadrid9, ReportWriterDocx
from apps.search.dataclasses import InidCode
from apps.bulletin.models import ClListOfficialBulletinsIp

from docx import Document

from pathlib import Path
import tempfile


class ReportItemDocxTmTestCase(TestCase):
    def setUp(self) -> None:
        self.document = Document()

    def test_inid_111(self):
        """Тестирует корректность добавления информации о номере регистрации."""
        # Номер регистрации есть в данных и разрешён для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    'RegistrationNumber': '111111'
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '111', 'Номер свідоцтва', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertTrue('(111)\tНомер свідоцтва: 111111' in p.text)

        # Номер регистрации есть в данных и не разрешён для отображения
        inid_data = [
            InidCode(4, '111', 'Номер свідоцтва', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertFalse('(111)\tНомер свідоцтва: 111111' in p.text)

        # Номера регистрации нет данных, разрешён для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '111', 'Номер свідоцтва', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertFalse('(111)\tНомер свідоцтва: 111111' in p.text)

        # Номер регистрации есть в данных, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    'RegistrationNumber': '111111'
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertFalse('(111)\tНомер свідоцтва: 111111' in p.text)

    def test_inid_540(self):
        """
        Тестирует добавление изображения.
        TODO: написать тест
        """
        pass

    def test_inid_141(self):
        """Тестирует корректность добавления информации о значении ИНИД
         (141) Дата закінчення строку дії реєстрації знака."""
        # Данные существуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "TerminationDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '141', 'Дата закінчення строку дії реєстрації знака', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn('(141)\tДата закінчення строку дії реєстрації знака: 23.02.2023', p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '141', 'Дата закінчення строку дії реєстрації знака', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(141)\tДата закінчення строку дії реєстрації знака: 23.02.2023', p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '141', 'Дата закінчення строку дії реєстрації знака', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(141)\tДата закінчення строку дії реєстрації знака: 23.02.2023', p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "TerminationDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(141)\tДата закінчення строку дії реєстрації знака: 23.02.2023', p.text)

    def test_inid_151(self):
        """Тестирует корректность добавления информации о значении ИНИД
        (151) Дата реєстрації."""
        # Данные существуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "RegistrationDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '151', 'Дата реєстрації', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn('(151)\tДата реєстрації: 23.02.2023', p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '151', 'Дата реєстрації', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(151)\tДата реєстрації: 23.02.2023', p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '151', 'Дата реєстрації', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(151)\tДата реєстрації: 23.02.2023', p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "RegistrationDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(151)\tДата реєстрації: 23.02.2023', p.text)

    def test_inid_181(self):
        """Тестирует корректность добавления информации о значении ИНИД
        (181) Очікувана дата закінчення строку дії реєстрації."""
        # Данные существуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ExpiryDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '181', 'Очікувана дата закінчення строку дії реєстрації', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn('(181)\tОчікувана дата закінчення строку дії реєстрації: 23.02.2023', p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '181', 'Очікувана дата закінчення строку дії реєстрації', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(181)\tОчікувана дата закінчення строку дії реєстрації: 23.02.2023', p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '181', 'Очікувана дата закінчення строку дії реєстрації', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(181)\tОчікувана дата закінчення строку дії реєстрації: 23.02.2023', p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ExpiryDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(181)\tОчікувана дата закінчення строку дії реєстрації: 23.02.2023', p.text)

    def test_inid_186(self):
        """Тестирует корректность добавления информации о значении ИНИД
        (186) Очікувана дата продовження строку дії реєстрації."""
        # Данные существуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ProlonagationExpiryDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '186', 'Очікувана дата продовження строку дії реєстрації', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn('(186)\tОчікувана дата продовження строку дії реєстрації: 23.02.2023', p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '186', 'Очікувана дата продовження строку дії реєстрації', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(186)\tОчікувана дата продовження строку дії реєстрації: 23.02.2023', p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '186', 'Очікувана дата продовження строку дії реєстрації', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(186)\tОчікувана дата продовження строку дії реєстрації: 23.02.2023', p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ProlonagationExpiryDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(186)\tОчікувана дата продовження строку дії реєстрації: 23.02.2023', p.text)

    def test_inid_210(self):
        """Тестирует корректность добавления информации о значении ИНИД
        (210) Порядковий номер заявки."""
        # Данные существуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ApplicationNumber": "m202200001"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '210', 'Порядковий номер заявки', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn('(210)\tПорядковий номер заявки: m202200001', p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '210', 'Порядковий номер заявки', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(210)\tПорядковий номер заявки: m202200001', p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '210', 'Порядковий номер заявки', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(210)\tПорядковий номер заявки: m202200001', p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ApplicationNumber": "m202200001"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(210)\tПорядковий номер заявки: m202200001', p.text)

    def test_inid_220(self):
        """Тестирует корректность добавления информации о значении ИНИД
        (220) Дата подання заявки."""
        # Данные существуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ApplicationDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '220', 'Дата подання заявки', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn('(220)\tДата подання заявки: 23.02.2023', p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '220', 'Дата подання заявки', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(220)\tДата подання заявки: 23.02.2023', p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '220', 'Дата подання заявки', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(220)\tДата подання заявки: 23.02.2023', p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ApplicationDate": "2023-02-23"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(220)\tДата подання заявки: 23.02.2023', p.text)

    def test_inid_531(self):
        """Тестирует корректность добавления информации о значении ИНИД
        (531) Віденська класифікація."""
        # Данные существуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "MarkImageDetails": {
                        "MarkImage": {
                            "MarkImageCategory": {
                                "CategoryCodeDetails": {
                                    "CategoryCode": [
                                        "01.01.01",
                                        "02.02.02",
                                    ]
                                }
                            }
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '531', 'Віденська класифікація', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn('(531)\tВіденська класифікація:\n01.01.01\n02.02.02', p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '531', 'Віденська класифікація', 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(531)\tВіденська класифікація:\n01.01.01\n02.02.02', p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '531', 'Віденська класифікація', 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(531)\tВіденська класифікація:\n01.01.01\n02.02.02', p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "MarkImageDetails": {
                        "MarkImage": {
                            "MarkImageCategory": {
                                "CategoryCodeDetails": {
                                    "CategoryCode": [
                                        "01.01.01",
                                        "02.02.02",
                                    ]
                                }
                            }
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn('(531)\tВіденська класифікація:\n01.01.01\n02.02.02', p.text)

    def test_inid_300(self):
        """Тестирует корректность добавления информации о значении ИНИД
        (300) Дані щодо пріоритету відповідно до Паризької конвенції та інші дані,
        пов'язані зі старшинством або реєстрацією знака у країні походження."""
        # Данные существуют и разрешены для отображения
        inid_title = "Дані щодо пріоритету відповідно до Паризької конвенції та інші дані, " \
                     "пов'язані зі старшинством або реєстрацією знака у країні походження"
        expected_str = f"(300)\t{inid_title}:\nm202300001; 2023-02-24; EN\nm202300001; 2023-02-25; US"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "PriorityDetails": {
                        "Priority": [
                            {
                                "PriorityCountryCode": "EN",
                                "PriorityDate": "2023-02-24",
                                "PriorityNumber": "m202300001"
                            },
                            {
                                "PriorityCountryCode": "US",
                                "PriorityDate": "2023-02-25",
                                "PriorityNumber": "m202300001"
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '300', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '300', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '300', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "PriorityDetails": {
                        "Priority": [
                            {
                                "PriorityCountryCode": "EN",
                                "PriorityDate": "2023-02-24",
                                "PriorityNumber": "m202300001"
                            },
                            {
                                "PriorityCountryCode": "US",
                                "PriorityDate": "2023-02-25",
                                "PriorityNumber": "m202300001"
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_731(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (731) Ім'я та адреса заявника
        TODO: Несколько заявителей
        """
        # Данные существуют и разрешены для отображения (заявка)
        inid_title = "Ім'я та адреса заявника"
        expected_str = f"(731)\t{inid_title}:\nІваненко Іван Іванович"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ApplicantDetails": {
                        "Applicant": [
                            {
                                "ApplicantAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "Адреса"
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 1
            }
        }
        inid_data = [
            InidCode(4, '731', inid_title, 1, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '731', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют и разрешены для отображения (свидетельство)
        expected_str = f"(731)\t{inid_title}:\nІваненко Іван Іванович\nАдреса (UA)"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ApplicantDetails": {
                        "Applicant": [
                            {
                                "ApplicantAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "Адреса"
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '731', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '731', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '731', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "ApplicantDetails": {
                        "Applicant": [
                            {
                                "ApplicantAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "Адреса"
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_732(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (732) Ім'я та адреса володільця реєстрації
        TODO: Несколько владельцев
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Ім'я та адреса володільця реєстрації"
        expected_str = f"(732)\t{inid_title}:\nІваненко Іван Іванович\nАдреса (UA)"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "HolderDetails": {
                        "Holder": [
                            {
                                "HolderAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "Адреса",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                },
                                            }
                                        }
                                    }
                                },
                            }
                        ]
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '732', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '732', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '732', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "HolderDetails": {
                        "Holder": [
                            {
                                "HolderAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "Адреса",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                },
                                            }
                                        }
                                    }
                                },
                            }
                        ]
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_740(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (740) Ім'я та адреса представника
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Ім'я та адреса представника"
        expected_str = f"(740)\t{inid_title}:\nІваненко Іван Іванович\nАдреса (UA)"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "RepresentativeDetails": {
                        "Representative": [
                            {
                                "RepresentativeAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "Адреса"
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameDetails": {
                                                        "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                    }
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '740', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '740', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '740', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "RepresentativeDetails": {
                        "Representative": [
                            {
                                "RepresentativeAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "Адреса"
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameDetails": {
                                                        "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                    }
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_750(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (750) Адреса для листування
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Ім'я та адреса представника"
        expected_str = f"(750)\t{inid_title}:\nІваненко Іван Іванович\nАдреса (UA)"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "CorrespondenceAddress": {
                        "CorrespondenceAddressBook": {
                            "Address": {
                                "AddressCountryCode": "UA",
                                "FreeFormatAddressLine": "Адреса"
                            },
                            "Name": {
                                "FreeFormatNameLine": "Іваненко Іван Іванович"
                            }
                        }
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '750', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '750', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '750', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "CorrespondenceAddress": {
                        "CorrespondenceAddressBook": {
                            "Address": {
                                "AddressCountryCode": "UA",
                                "FreeFormatAddressLine": "Адреса"
                            },
                            "Name": {
                                "FreeFormatNameLine": "Іваненко Іван Іванович"
                            }
                        }
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_591(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (591) Інформація щодо заявлених кольорів
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Інформація щодо заявлених кольорів"
        expected_str = f"(591)\t{inid_title}: чорний, білий"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "MarkImageDetails": {
                        "MarkImage": {
                            "MarkImageColourClaimedText": [
                                {
                                    "#text": "чорний"
                                },
                                {
                                    "#text": "білий"
                                }
                            ]
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '591', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '591', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '591', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "MarkImageDetails": {
                        "MarkImage": {
                            "MarkImageColourClaimedText": [
                                {
                                    "#text": "чорний"
                                },
                                {
                                    "#text": "білий"
                                }
                            ]
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_511(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (511) Індекси Ніццької класифікації
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Індекси Ніццької класифікації"
        expected_str = f"(511)\t{inid_title}: 1, 2"
        expected_str += f"\nКл. 1:\tзначення 1.1; значення 1.2"
        expected_str += f"\nКл. 2:\tзначення 2.1; значення 2.2"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "GoodsServicesDetails": {
                        "GoodsServices": {
                            "ClassDescriptionDetails": {
                                "ClassDescription": [
                                    {
                                        "ClassNumber": 1,
                                        "ClassificationTermDetails": {
                                            "ClassificationTerm": [
                                                {
                                                    "ClassificationTermText": "значення 1.1"
                                                },
                                                {
                                                    "ClassificationTermText": "значення 1.2"
                                                }
                                            ]
                                        }
                                    },
                                    {
                                        "ClassNumber": 2,
                                        "ClassificationTermDetails": {
                                            "ClassificationTerm": [
                                                {
                                                    "ClassificationTermText": "значення 2.1"
                                                },
                                                {
                                                    "ClassificationTermText": "значення 2.2"
                                                }
                                            ]
                                        }
                                    }
                                ]
                            }
                        }
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '511', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '511', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '511', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '511', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "GoodsServicesDetails": {
                        "GoodsServices": {
                            "ClassDescriptionDetails": {
                                "ClassDescription": [
                                    {
                                        "ClassNumber": 1,
                                        "ClassificationTermDetails": {
                                            "ClassificationTerm": [
                                                {
                                                    "ClassificationTermText": "значення 1.1"
                                                },
                                                {
                                                    "ClassificationTermText": "значення 1.2"
                                                }
                                            ]
                                        }
                                    },
                                    {
                                        "ClassNumber": 2,
                                        "ClassificationTermDetails": {
                                            "ClassificationTerm": [
                                                {
                                                    "ClassificationTermText": "значення 2.1"
                                                },
                                                {
                                                    "ClassificationTermText": "значення 2.2"
                                                }
                                            ]
                                        }
                                    }
                                ]
                            }
                        }
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_app_input_date(self):
        """
        Тестирует корректность добавления информации о значении поля "Дата надходження матеріалів заявки до НОІВ"
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Дата надходження матеріалів заявки до НОІВ"
        expected_str = f"{inid_title}: 25.02.2023"
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "app_input_date": "2023-02-25",
                }
            },
            'search_data': {
                'obj_state': 1
            }
        }
        inid_data = [
            InidCode(4, '221', inid_title, 1, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '221', inid_title, 1, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '221', inid_title, 1, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "app_input_date": "2023-02-25",
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_441(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (441) Дата публікації відомостей про заявку та номер бюлетня
        """
        inid_title = "Дата публікації відомостей про заявку та номер бюлетня"
        expected_str_ua = f"(441)\t{inid_title}: 26.02.2023, бюл. №1"
        expected_str_en = f"(441)\t{inid_title}: 26.02.2023, bul. №1"

        # Данные существуют и разрешены для отображения, номер бюллетеня присутствует в данных JSON
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "Code_441": "2023-02-26",
                    "Code_441_BulNumber": "1"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '441', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data, 'ua')
        p = item.write(self.document)
        self.assertIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertIn(expected_str_en, p.text)

        # Данные существуют и разрешены для отображения, номер бюлетня отсутствует в данных JSON
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "Code_441": "2023-02-26"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '441', inid_title, 2, True)
        ]
        ClListOfficialBulletinsIp.objects.create(
            bul_number=1,
            bul_date='2023-02-26',
            date_from='2023-02-20',
            date_to='2023-02-28',
        )
        item = ReportItemDocxTM(biblio_data, inid_data, 'ua')
        p = item.write(self.document)
        self.assertIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertIn(expected_str_en, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '441', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '441', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "Code_441": "2023-02-26"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

    def test_inid_450(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (450) Дата публікації відомостей про видачу свідоцтва
        """
        inid_title = "Дата публікації відомостей про видачу свідоцтва"
        expected_str_ua = f"(450)\t{inid_title}: 27.02.2023, бюл. № 1/2023"
        expected_str_en = f"(450)\t{inid_title}: 27.02.2023, bul. № 1/2023"

        # Данные существуют и разрешены для отображения, номер бюллетеня присутствует в данных JSON
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "PublicationDetails": [
                        {
                            "PublicationDate": "2023-02-27",
                            "PublicationIdentifier": "1/2023"
                        }
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '450', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data, 'ua')
        p = item.write(self.document)
        self.assertIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertIn(expected_str_en, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(4, '450', inid_title, 2, False)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '450', inid_title, 2, True)
        ]
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    "PublicationDetails": [
                        {
                            "PublicationDate": "2023-02-27",
                            "PublicationIdentifier": "1/2023"
                        }
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxTM(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxTM(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)


class ReportItemDocxIDTestCase(TestCase):
    def setUp(self) -> None:
        self.document = Document()

    def test_inid_21(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (21) Номер заявки
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Номер заявки"
        expected_str = f"(21)\t{inid_title}: s202300001"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignApplicationNumber": "s202300001"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '21', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '21', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '21', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignApplicationNumber": "s202300001"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_51(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (51) Індекс(и) Міжнародної класифікації промислових зразків (МКПЗ)
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Індекс(и) Міжнародної класифікації промислових зразків (МКПЗ)"
        expected_str = f"(51)\t{inid_title}: 01-00, 02-00"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "IndicationProductDetails": [
                        {
                            "Class": "01-00"
                        },
                        {
                            "Class": "02-00"
                        },
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '51', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '51', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '51', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "IndicationProductDetails": [
                        {
                            "Class": "01-00"
                        },
                        {
                            "Class": "02-00"
                        },
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_11(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (11) Номер реєстрації (патенту)
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Номер реєстрації (патенту)"
        expected_str = f"(11)\t{inid_title}: 11111"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "RegistrationNumber": "11111"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '11', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '11', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '11', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "RegistrationNumber": "11111"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_22(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (22) Дата подання заявки
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Дата подання заявки"
        expected_str = f"(22)\t{inid_title}: 26.02.2023"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignApplicationDate": "2023-02-26"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '22', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '22', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '22', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignApplicationDate": "2023-02-26"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_23(self):
        """
        Тестирует корректность добавления информации о значении ИНИД (23) Дата виставкового пріоритету
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Дата виставкового пріоритету"
        expected_str = f"(23)\t{inid_title}:\n27.02.2023; UA\n28.02.2023; UA"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "ExhibitionPriorityDetails": [
                        {
                            "ExhibitionDate": "2023-02-27",
                            "ExhibitionCountryCode": "UA"
                        },
                        {
                            "ExhibitionDate": "2023-02-28",
                            "ExhibitionCountryCode": "UA"
                        }
                    ],
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '23', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '23', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '23', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "ExhibitionPriorityDetails": [
                        {
                            "ExhibitionDate": "2023-02-27",
                            "ExhibitionCountryCode": "UA"
                        },
                        {
                            "ExhibitionDate": "2023-02-28",
                            "ExhibitionCountryCode": "UA"
                        }
                    ],
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_24(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (24) Дата, з якої є чинними права на промисловий зразок
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Дата, з якої є чинними права на промисловий зразок"
        expected_str = f"(24)\t{inid_title}: 27.02.2023"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "RecordEffectiveDate": "2023-02-27"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '24', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '24', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '24', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    'DesignDetails': {
                        "RecordEffectiveDate": "2023-02-27"
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_28(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (28) Кількість заявлених варіантів
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Кількість заявлених варіантів"
        expected_str = f"(28)\t{inid_title}: 2"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "TotalSpecimen": 2
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '28', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '28', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '28', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    'DesignDetails': {
                        "TotalSpecimen": 2
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_31(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (31) Номер попередньої заявки відповідно до Паризької конвенції
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Номер попередньої заявки відповідно до Паризької конвенції"
        expected_str = f"(31)\t{inid_title}: s202300001"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "PriorityDetails": [
                        {
                            "PriorityNumber": "s202300001",
                            "PriorityDate": "2023-02-27",
                            "PriorityCountryCode": "UA"
                        }
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '31', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Альтернативная структура данных (есть ключ Priority)
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "PriorityDetails": {
                        'Priority': [
                            {
                                "PriorityNumber": "s202300001",
                                "PriorityDate": "2023-02-27",
                                "PriorityCountryCode": "UA"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '31', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '31', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    'DesignDetails': {
                        "PriorityDetails": [
                            {
                                "PriorityNumber": "s202300001",
                                "PriorityDate": "2023-02-27",
                                "PriorityCountryCode": "UA"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_32(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (32) Номер попередньої заявки відповідно до Паризької конвенції
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Дата подання попередньої заявки відповідно до Паризької конвенції"
        expected_str = f"(32)\t{inid_title}: 27.02.2023"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "PriorityDetails": [
                        {
                            "PriorityNumber": "s202300001",
                            "PriorityDate": "2023-02-27",
                            "PriorityCountryCode": "UA"
                        }
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '32', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Альтернативная структура данных (есть ключ Priority)
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "PriorityDetails": {
                        'Priority': [
                            {
                                "PriorityNumber": "s202300001",
                                "PriorityDate": "2023-02-27",
                                "PriorityCountryCode": "UA"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '32', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '32', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    'DesignDetails': {
                        "PriorityDetails": [
                            {
                                "PriorityNumber": "s202300001",
                                "PriorityDate": "2023-02-27",
                                "PriorityCountryCode": "UA"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_33(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (33) Двобуквений код держави-учасниці Паризької конвенції, до якої подано попередню заявку
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Дата подання попередньої заявки відповідно до Паризької конвенції"
        expected_str = f"(33)\t{inid_title}: UA"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "PriorityDetails": [
                        {
                            "PriorityNumber": "s202300001",
                            "PriorityDate": "2023-02-27",
                            "PriorityCountryCode": "UA"
                        }
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '33', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Альтернативная структура данных (есть ключ Priority)
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "PriorityDetails": {
                        'Priority': [
                            {
                                "PriorityNumber": "s202300001",
                                "PriorityDate": "2023-02-27",
                                "PriorityCountryCode": "UA"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '33', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '33', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    'DesignDetails': {
                        "PriorityDetails": [
                            {
                                "PriorityNumber": "s202300001",
                                "PriorityDate": "2023-02-27",
                                "PriorityCountryCode": "UA"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_45(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (45) Дата публікації відомостей про видачу патенту та номер бюлетеня
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Дата публікації відомостей про заявку та номер бюлетня"
        expected_str_ua = f"(45)\t{inid_title}: 27.02.2023, бюл. № 1/2023"
        expected_str_en = f"(45)\t{inid_title}: 27.02.2023, bul. № 1/2023"

        # Данные существуют и разрешены для отображения, номер бюллетеня присутствует в данных JSON
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "RecordPublicationDetails": [
                        {
                            "PublicationIdentifier": "1/2023",
                            "PublicationDate": "2023-02-27"
                        }
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '45', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data, 'ua')
        p = item.write(self.document)
        self.assertIn(expected_str_ua, p.text)
        item = ReportItemDocxID(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertIn(expected_str_en, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '45', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxID(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '45', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxID(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "RecordPublicationDetails": [
                        {
                            "PublicationIdentifier": "1/2023",
                            "PublicationDate": "2023-02-27"
                        }
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxID(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

    def test_inid_54(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (54) Назва промислового зразка
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Назва промислового зразка"
        expected_str = f"(54)\t{inid_title}: Назва промислового зразка"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignTitle": "Назва промислового зразка",
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '54', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '54', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '54', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    'DesignDetails': {
                        "DesignTitle": "Назва промислового зразка",
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_55(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (55) Зображення промислового зразка та вказівка відносно кольорів
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Зображення промислового зразка та вказівка відносно кольорів"
        expected_str = f"(55)\t{inid_title}:\n1-й варіант - чорний, білий\n2-й варіант - червоний, зелений"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignSpecimenDetails": [
                        {
                            "Colors": {
                                "Color": "чорний, білий"
                            }
                        },
                        {
                            "Colors": {
                                "Color": "червоний, зелений"
                            }
                        },
                    ]
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '55', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '55', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '55', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    'DesignDetails': {
                        "DesignTitle": "Назва промислового зразка",
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_71(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (71) Ім'я (імена) та адреса (адреси) заявника (заявників)
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Ім'я (імена) та адреса (адреси) заявника (заявників)"
        expected_str = f"(71)\t{inid_title}:\nІваненко Іван Іванович\nм.Київ\nПетренко Петро Петрович\nм.Дніпро"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "ApplicantDetails": {
                        "Applicant": [
                            {
                                "ApplicantAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Київ",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                            {
                                "ApplicantAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Дніпро",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Петренко Петро Петрович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '71', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '71', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '71', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "ApplicantDetails": {
                        "Applicant": [
                            {
                                "ApplicantAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Київ",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                            {
                                "ApplicantAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Дніпро",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Петренко Петро Петрович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_72(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (72) Ім'я (імена) автора (авторів), якщо воно відоме
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Ім'я (імена) автора (авторів), якщо воно відоме"
        expected_str = f"(72)\t{inid_title}:\nІваненко Іван Іванович\nПетренко Петро Петрович"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignerDetails": {
                        "Designer": [
                            {
                                "DesignerAddressBook": {
                                    "FormattedNameAddress": {
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                },
                                            }
                                        }
                                    }
                                },
                                "Publicated": True
                            },
                            {
                                "DesignerAddressBook": {
                                    "FormattedNameAddress": {
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Вікторов Віктор Вікторович"
                                                },
                                            }
                                        }
                                    }
                                },
                                "Publicated": False
                            },
                            {
                                "DesignerAddressBook": {
                                    "FormattedNameAddress": {
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Петренко Петро Петрович"
                                                },
                                            }
                                        }
                                    }
                                },
                                "Publicated": True
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '72', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '72', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '72', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "DesignerDetails": {
                        "Designer": [
                            {
                                "DesignerAddressBook": {
                                    "FormattedNameAddress": {
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                },
                                            }
                                        }
                                    }
                                },
                                "Publicated": True
                            },
                            {
                                "DesignerAddressBook": {
                                    "FormattedNameAddress": {
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Вікторов Віктор Вікторович"
                                                },
                                            }
                                        }
                                    }
                                },
                                "Publicated": False
                            },
                            {
                                "DesignerAddressBook": {
                                    "FormattedNameAddress": {
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Петренко Петро Петрович"
                                                },
                                            }
                                        }
                                    }
                                },
                                "Publicated": True
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_73(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (73) Ім’я або повне найменування власника(ів) патенту, його адреса та двобуквений код держави
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Ім’я або повне найменування власника(ів) патенту, його адреса та двобуквений код держави"
        expected_str = f"(73)\t{inid_title}:\nІваненко Іван Іванович\nм.Київ\n(UA)\nПетренко Петро Петрович\nм.Дніпро\n(EN)"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "HolderDetails": {
                        "Holder": [
                            {
                                "HolderAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Київ",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                            {
                                "HolderAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "EN",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Дніпро",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Петренко Петро Петрович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '73', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '73', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '73', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "HolderDetails": {
                        "Holder": [
                            {
                                "HolderAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "UA",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Київ",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                            {
                                "HolderAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "AddressCountryCode": "EN",
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Дніпро",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Петренко Петро Петрович"
                                                }
                                            }
                                        }
                                    }
                                }
                            },
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_74(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (74) Ім'я (імена) та адреса (адреси) представника (представників)
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Ім'я (імена) та адреса (адреси) представника (представників)"
        expected_str = f"(74)\t{inid_title}:\nІваненко Іван Іванович\nм.Київ"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "RepresentativeDetails": {
                        "Representative": [
                            {
                                "RepresentativeAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Київ",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '74', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '74', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '74', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "RepresentativeDetails": {
                        "Representative": [
                            {
                                "RepresentativeAddressBook": {
                                    "FormattedNameAddress": {
                                        "Address": {
                                            "FreeFormatAddress": {
                                                "FreeFormatAddressLine": "м.Київ",
                                            }
                                        },
                                        "Name": {
                                            "FreeFormatName": {
                                                "FreeFormatNameDetails": {
                                                    "FreeFormatNameLine": "Іваненко Іван Іванович"
                                                }
                                            }
                                        }
                                    }
                                }
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_98(self):
        """
        Тестирует корректность добавления информации о значении
        ИНИД (98) Адреса для листування
        """
        # Данные существуют и разрешены для отображения
        inid_title = "Адреса для листування"
        expected_str = f"(98)\t{inid_title}:\nІваненко Іван Іванович\nм.Київ\n(UA)"
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "CorrespondenceAddress": {
                        "CorrespondenceAddressBook": {
                            "FormattedNameAddress": {
                                "Address": {
                                    "AddressCountryCode": "UA",
                                    "FreeFormatAddress": {
                                        "FreeFormatAddressLine": "м.Київ"
                                    }
                                },
                                "Name": {
                                    "FreeFormatName": {
                                        "FreeFormatNameDetails": {
                                            "FreeFormatNameLine": "Іваненко Іван Іванович"
                                        }
                                    }
                                }
                            }
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '98', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные существуют и запрещены для отображения
        inid_data = [
            InidCode(6, '98', inid_title, 2, False)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'Design': {
                'DesignDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(6, '98', inid_title, 2, True)
        ]
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные существуют, ИНИД код - нет
        biblio_data = {
            'Design': {
                'DesignDetails': {
                    "CorrespondenceAddress": {
                        "CorrespondenceAddressBook": {
                            "FormattedNameAddress": {
                                "Address": {
                                    "AddressCountryCode": "UA",
                                    "FreeFormatAddress": {
                                        "FreeFormatAddressLine": "м.Київ"
                                    }
                                },
                                "Name": {
                                    "FreeFormatName": {
                                        "FreeFormatNameDetails": {
                                            "FreeFormatNameLine": "Іваненко Іван Іванович"
                                        }
                                    }
                                }
                            }
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxID(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)


class ReportItemDocxMadridTestCase(TestCase):
    def setUp(self) -> None:
        self.document = Document()

    def test_inid_540(self):
        """Тестирует корректность добавления информации об
        ИНИД (540) Зображення торговельної марки."""
        # Разрешено для отображения
        inid_title = "Зображення торговельної марки"
        expected_str = f"(540)\t{inid_title}:"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '540', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Запрещено для отображения
        inid_data = [
            InidCode(9, '540', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_111(self):
        """Тестирует корректность добавления информации об
        ИНИД (111) Номер міжнародної реєстрації."""
        # Номер регистрации есть в данных и разрешён для отображения
        inid_title = "Номер міжнародної реєстрації"
        expected_str = f"(111)\t{inid_title}: 111111"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    '@INTREGN': '111111'
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '111', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Номер регистрации есть в данных и не разрешён для отображения
        inid_data = [
            InidCode(9, '111', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Номера регистрации нет данных, разрешён для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '111', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Номер регистрации есть в данных, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    '@INTREGN': '111111'
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_151(self):
        """Тестирует корректность добавления информации об
        ИНИД (151) Дата міжнародної реєстрації."""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Дата міжнародної реєстрації"
        expected_str = f"(151)\t{inid_title}: 01.03.2023"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "@INTREGD": "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '151', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '151', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '151', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "@INTREGD" : "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_180(self):
        """Тестирует корректность добавления информации об
        ИНИД (180) Очікувана дата закінчення строку дії реєстрації/продовження."""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Очікувана дата закінчення строку дії реєстрації/продовження"
        expected_str = f"(180)\t{inid_title}: 01.03.2023"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "@INTREGD": "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '180', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '180', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '180', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "@INTREGD": "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_441(self):
        """Тестирует корректность добавления информации об
           ИНИД (441) Дата публікації відомостей про міжнародну реєстрацію торговельної марки,
           що надійшла для проведення експертизи."""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Дата публікації відомостей про міжнародну реєстрацію торговельної марки, " \
                     "що надійшла для проведення експертизи"
        expected_str_ua = f"(441)\t{inid_title}: 01.03.2023, бюл. № 8"
        expected_str_en = f"(441)\t{inid_title}: 01.03.2023, bul. № 8"
        ClListOfficialBulletinsIp.objects.create(
            bul_number=8,
            bul_date='2023-03-01',
            date_from='2023-02-20',
            date_to='2023-03-28',
        )
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "Code_441": "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '441', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertIn(expected_str_en, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '441', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '441', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "Code_441": "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

    def test_inid_450(self):
        """Тестирует корректность добавления информации об
        ИНИД (450) Дата публікації відомостей про міжнародну реєстрацію та номер бюлетеню Міжнародного бюро ВОІВ"""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Дата публікації відомостей про міжнародну реєстрацію та номер бюлетеню Міжнародного бюро ВОІВ"
        expected_str_ua = f"(450)\t{inid_title}: 01.03.2023, бюл. № 2023/8 Gaz"
        expected_str_en = f"(450)\t{inid_title}: 01.03.2023, bul. № 2023/8 Gaz"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "ENN": {
                        "@PUBDATE": "2023-03-01",
                        "@GAZNO": "2023/8 Gaz"
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '450', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertIn(expected_str_en, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '450', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '450', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "ENN": {
                        "@PUBDATE": "2023-03-01",
                        "@GAZNO": "2023/8 Gaz"
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str_ua, p.text)
        item = ReportItemDocxMadrid9(biblio_data, inid_data, 'en')
        p = item.write(self.document)
        self.assertNotIn(expected_str_en, p.text)

    def test_inid_732(self):
        """Тестирует корректность добавления информации об
        ИНИД (732) Ім'я та адреса володільця реєстрації."""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Ім'я та адреса володільця реєстрації"
        expected_str = f"(732)\t{inid_title}:\nName\nAddr. str. 1\nAddr. str. 2\n(UA)"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "HOLGR": {
                        "NAME": {
                            "NAMEL": "Name"
                        },
                        "ADDRESS": {
                            "ADDRL": [
                                "Addr. str. 1",
                                "Addr. str. 2"
                            ],
                            "COUNTRY": "UA"
                        },
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '732', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '732', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '732', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "HOLGR": {
                        "NAME": {
                            "NAMEL": "Name"
                        },
                        "ADDRESS": {
                            "ADDRL": [
                                "Addr. str. 1",
                                "Addr. str. 2"
                            ],
                            "COUNTRY": "UA"
                        },
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_740(self):
        """Тестирует корректность добавления информации об
        ИНИД (740) Ім'я та адреса представника."""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Ім'я та адреса представника"
        expected_str = f"(740)\t{inid_title}:\nName\nAddr. str. 1\nAddr. str. 2\n(UA)"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "REPGR": {
                        "NAME": {
                            "NAMEL": "Name"
                        },
                        "ADDRESS": {
                            "ADDRL": [
                                "Addr. str. 1",
                                "Addr. str. 2",
                            ],
                            "COUNTRY": "UA"
                        }
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '740', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '740', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '740', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "REPGR": {
                        "NAME": {
                            "NAMEL": "Name"
                        },
                        "ADDRESS": {
                            "ADDRL": [
                                "Addr. str. 1",
                                "Addr. str. 2",
                            ],
                            "COUNTRY": "UA"
                        }
                    },
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_821(self):
        """Тестирует корректность добавления информации об
        ИНИД (821) Базова заявка."""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Базова заявка"
        expected_str = f"(821)\t{inid_title}: 01.03.2023, 12345"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "BASGR": {
                        "BASAPPGR": {
                            "BASAPPD": "2023-03-01",
                            "BASAPPN": "12345"
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '821', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '821', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '821', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "BASGR": {
                        "BASAPPGR": {
                            "BASAPPD": "2023-03-01",
                            "BASAPPN": "12345"
                        }
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_891(self):
        """Тестирует корректность добавления информации об
        ИНИД (891) Дата територіального поширення міжнародної реєстрації."""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Дата територіального поширення міжнародної реєстрації"
        expected_str = f"(891)\t{inid_title}: 01.03.2023"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "@REGEDAT": "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '891', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '891', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '891', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "@REGEDAT": "2023-03-01"
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

    def test_inid_511(self):
        """Тестирует корректность добавления информации об
        ИНИД (511) Індекс (індекси) МКТП та перелік товарів і послуг"""
        # Данные присутствуют и разрешены для отображения
        inid_title = "Індекс (індекси) МКТП та перелік товарів і послуг"
        expected_str = f"(511)\t{inid_title}:\nКл. 01: qwe, rty, uio\nКл. 02: asd, fgh, jkl"
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "BASICGS": {
                        "@NICEVER": "10",
                        "GSGR": [
                            {
                                "@NICCLAI": "01",
                                "GSTERMEN": "qwe, rty, uio"
                            },
                            {
                                "@NICCLAI": "02",
                                "GSTERMEN": "asd, fgh, jkl"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '511', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertIn(expected_str, p.text)

        # Данные присутствуют и не разрешены для отображения
        inid_data = [
            InidCode(9, '511', inid_title, 2, False)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные отсутствуют и разрешены для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {}
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(9, '511', inid_title, 2, True)
        ]
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)

        # Данные присутствуют, ИНИД-кода нет в данных для отображения
        biblio_data = {
            'MadridTradeMark': {
                'TradeMarkDetails': {
                    "BASICGS": {
                        "@NICEVER": "10",
                        "GSGR": [
                            {
                                "@NICCLAI": "01",
                                "GSTERMEN": "qwe, rty, uio"
                            },
                            {
                                "@NICCLAI": "02",
                                "GSTERMEN": "asd, fgh, jkl"
                            }
                        ]
                    }
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = []
        item = ReportItemDocxMadrid9(biblio_data, inid_data)
        p = item.write(self.document)
        self.assertNotIn(expected_str, p.text)


class ReportWriterDocxTestCase(TestCase):
    """Тестирует класс создания файла отчёта в формате .docx."""
    report_path = Path(tempfile.gettempdir()) / 'report.docx'
    # report_path = Path('report.docx')

    def test_generates_report(self):
        """Тестирует создание отчёта."""
        biblio_data = {
            'TradeMark': {
                'TrademarkDetails': {
                    'RegistrationNumber': '111111'
                }
            },
            'search_data': {
                'obj_state': 2
            }
        }
        inid_data = [
            InidCode(4, '111', 'Номер свідоцтва', 2, True)
        ]
        items = [ReportItemDocxTM(biblio_data, inid_data)]

        file_path = Path(tempfile.gettempdir()) / 'report.docx'
        writer = ReportWriterDocx(items)
        writer.generate(file_path)
        self.assertTrue(file_path.exists())

from django.conf import settings
from django.http import HttpResponse
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.utils.translation import ugettext as _
from django.contrib.auth import get_user_model
from django.db.models import Prefetch
from elasticsearch import Elasticsearch
from elasticsearch_dsl import Search, Q, A
from elasticsearch_dsl.aggs import Terms, Nested
from .models import ObjType, InidCodeSchedule, OrderService, SortParameter, IpcAppList, PaidServicesSettings, IpcCode
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm
import re
import time
import datetime
import os
from uma.utils import iterable
import xlsxwriter
from apps.bulletin.models import ClListOfficialBulletinsIp


def get_search_groups(search_data):
    """Разбивка поисковых данных на поисковые группы"""
    search_groups = []
    for obj_type in ObjType.objects.all():
        # Поисковые запросы на заявки
        search_groups.append({
            'obj_type': obj_type,
            'obj_state': 1,
            'search_params': list(filter(
                lambda x: str(obj_type.pk) in x['obj_type'] and '1' in x['obj_state'],
                search_data
            ))
        })
        # Поисковые запросы на охранные документы
        search_groups.append({
            'obj_type': obj_type,
            'obj_state': 2,
            'search_params': list(filter(
                lambda x: str(obj_type.pk) in x['obj_type'] and '2' in x['obj_state'],
                search_data
            ))
        })
    # Фильтрация пустых групп
    search_groups = filter(lambda x: len(x['search_params']) > 0, search_groups)
    return list(search_groups)


def prepare_query(query, field_type):
    """Обрабатывает строку расширенного запроса пользователя."""
    if field_type == 'date':
        # Форматирование дат
        query = re.sub(r'(\d{1,2})\.(\d{1,2})\.(\d{4})', '\\3-\\2-\\1', query)

        # На вход происходит получение диапазона дат в формате "дд.мм.гггг ~ дд.мм.гггг"
        # его необходимо преобразовать запрос ElasticSearch
        dates = query.split(' ~ ')
        try:
            query = f"[{dates[0]} TO {dates[1]}]"
        except IndexError:
            pass

    query = query.replace(
        " ТА ", " AND "
    ).replace(
        " АБО ", " OR "
    ).replace(
        " НЕ ", " NOT "
    ).replace(
        "/", "\\/"
    ).replace(
        ":", "\\:"
    )
    return query


def get_elastic_results(search_groups, user):
    """Поиск в ElasticSearch по группам."""
    qs_list = []
    for group in search_groups:
        if group['search_params']:
            # Идентификаторы schedule_type для заявок или охранных документов
            schedule_type_ids = (10, 11, 12, 13, 14, 15) if group['obj_state'] == 1 else (3, 4, 5, 6, 7, 8, 16, 17, 18, 19, 30, 32)
            qs = None

            for search_param in group['search_params']:
                # Поле поиска ElasticSearch
                inid_schedule = InidCodeSchedule.objects.filter(
                    ipc_code__id=search_param['ipc_code'],
                    schedule_type__id__in=schedule_type_ids
                ).first()

                # Проверка доступно ли поле для поиска
                if inid_schedule.enable_search and inid_schedule.elastic_index_field is not None:
                    query = prepare_query(search_param['value'], inid_schedule.elastic_index_field.field_type)

                    if inid_schedule.elastic_index_field.field_type == 'text':
                        if '*' in query:
                            fields = [
                                # f"{inid_schedule.elastic_index_field.field_name}^2",
                                f"{inid_schedule.elastic_index_field.field_name}.exact^2",
                            ]
                        else:
                            fields = [
                                # f"{inid_schedule.elastic_index_field.field_name}^2",
                                f"{inid_schedule.elastic_index_field.field_name}.exact^2",
                                f"{inid_schedule.elastic_index_field.field_name}.*",
                            ]
                    else:
                        fields = [
                            f"{inid_schedule.elastic_index_field.field_name}",
                        ]

                    q = Q(
                        'query_string',
                        query=query,
                        fields=fields,
                        quote_field_suffix=".exact",
                        default_operator='AND'
                    )

                    # Nested тип
                    if inid_schedule.elastic_index_field.parent:
                        q = Q(
                            'nested',
                            path=inid_schedule.elastic_index_field.parent.field_name,
                            query=q,
                        )

                    if not qs:
                        qs = q
                    else:
                        qs &= q

            if qs is not None:
                qs &= Q('query_string', query=f"{group['obj_type'].pk}", default_field='Document.idObjType')
                qs &= Q('query_string', query=f"{group['obj_state']}", default_field='search_data.obj_state')

                # Не включать в список результатов заявки, по которым выдан патент
                qs = filter_bad_apps(qs)

                qs_list.append(qs)

    # Формирование результирующего запроса
    qs_result = None
    for qs in qs_list:
        if qs_result is None:
            qs_result = qs
        else:
            qs_result |= qs
    client = Elasticsearch(settings.ELASTIC_HOST, timeout=settings.ELASTIC_TIMEOUT)
    s = Search(using=client, index=settings.ELASTIC_INDEX_NAME).query(
        qs_result
    ).source(
        excludes=["*.DocBarCode", "*.DOCBARCODE"]
    ).sort('_score')

    return s


def get_client_ip(request):
    """Возвращает IP-адрес пользователя."""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip


class ResultsProxy(object):
    """
    A proxy object for returning Elasticsearch results that is able to be
    passed to a Paginator.
    """
    def __init__(self, s):
        self.s = s

    def __len__(self):
        return self.s['total']

    def __getitem__(self, item):
        return self.s['items']


def paginate_results(s, page, paginate_by=10):
    """Пагинатор для результов запроса ElasticSearch"""
    paginator = Paginator(ResultsProxy(s), paginate_by)
    page_number = page
    try:
        page = paginator.page(page_number)
    except PageNotAnInteger:
        page = paginator.page(1)
    except EmptyPage:
        page = paginator.page(paginator.num_pages)
    return page


def filter_bad_apps(qs):
    """Исключает из результатов запроса заявки, которые не положено публиковать."""
    # Не показывать заявки, по которым выдан охранный документ
    qs &= ~Q('query_string', query="Document.Status:3 AND search_data.obj_state:1")
    qs &= ~Q('query_string', query="_exists_:Claim.I_11")
    # qs &= ~Q('query_string', query="Document.idObjType:10 OR Document.idObjType:13")
    # qs &= ~Q(
    #     'query_string',
    #     query='(Document.MarkCurrentStatusCodeType:{* TO 2000} AND search_data.app_date:{* TO 2020-07-18}) '
    #           'OR (search_data.obj_state:1 AND Document.idObjType:4 '
    #           'AND NOT _exists_:TradeMark.TrademarkDetails.Code_441 AND search_data.app_date:[2020-07-18 TO *})'
    # )

    # Не показывать охранные документы, у которых дата выдачи больше сегодняшней
    # qs &= ~Q(
    #     'query_string',
    #     query=f">{datetime.datetime.now().strftime('%Y-%m-%d')}",
    #     default_field='search_data.rights_date'
    # )

    return qs


def filter_unpublished_apps(user, qs):
    """Исключает из результатов запроса неопубликованные заявки для обычных пользователей."""
    if not user.is_anonymous and user.is_vip():
        # Если пользователь является членом группы "посадовці, чиновники",
        # то ему возвращаются все результаты без фильтров
        return qs

    """Фильтры для обычных пользователей."""
    # Не показывать заявки по полезным моделям
    # filter_qs = ~Q('query_string', query="search_data.obj_state:1 AND Document.idObjType:2")

    # paid_services_settings, created = PaidServicesSettings.objects.get_or_create()

    """ ВРЕМЕННО ОТКРЫТЬ ДОСТУП ВСЕМ """
    # if not paid_services_settings.enabled:
    #     # Не показывать заявки на знаки со статусом 1000
    #     filter_qs &= ~Q('query_string', query="Document.MarkCurrentStatusCodeType:1000")

    # Показывать только заявки с датой заяки (но показывать все КЗПТ и заявки на знаки для товаров и услуг с кодом 1000)
    # filter_qs = Q(
    #     'query_string',
    #     query="_exists_:search_data.app_date OR Document.idObjType:5 "
    #           "OR (NOT _exists_:search_data.app_date AND Document.MarkCurrentStatusCodeType:1000)"
    # )
    # Для заявок на изобретения нужно чтоб существовал I_43.D
    # filter_qs &= ~Q('query_string', query="NOT _exists_:Claim.I_43.D AND search_data.obj_state:1 AND Document.idObjType:1")
    # Для заявок на знаки для товаров и услуг нужно чтоб существовали платежи
    # filter_qs &= ~Q(
    #     'query_string',
    #     query="NOT _exists_:TradeMark.PaymentDetails "
    #           "AND search_data.obj_state:1 "
    #           "AND Document.idObjType:4 "
    #           "AND NOT Document.MarkCurrentStatusCodeType:1000"
    # )

    # try:
    #     # Для обычных пользователей, авторизированных по ЭЦП применяются все фильтры, но показываются "их" заявки
    #     user_name = user.certificateowner.pszSubjFullName.strip()
    #     qs &= filter_qs | Q('query_string', query=f"search_data.applicant:\"*{user_name}*\"") \
    #           | Q('query_string', query=f"search_data.inventor:\"*{user_name}*\"") \
    #           | Q('query_string', query=f"search_data.owner:\"*{user_name}*\"") \
    #           | Q('query_string', query=f"search_data.agent:\"*{user_name}*\"")
    # except (get_user_model().certificateowner.RelatedObjectDoesNotExist, AttributeError):
    #     # Для обычных пользователей без ЭЦП или анонимных пользователей применяются все фильтры
    #     qs &= filter_qs

    return qs


def sort_results(s, sort_by_value):
    """Добавляет параметр сортировки."""
    try:
        sort_param = SortParameter.objects.values(
            'ordering',
            'search_field__elastic_index_field__field_name',
            'search_field__elastic_index_field__field_type'
        ).get(value=sort_by_value, is_enabled=True)
    except SortParameter.DoesNotExist:
        pass
    else:
        param = sort_param['search_field__elastic_index_field__field_name'].replace('*', '')
        if sort_param['search_field__elastic_index_field__field_type'] == 'text':
            param = f"{param}.raw"
        s = s.sort(
            {
                param: {
                    "order": sort_param['ordering'],
                    "missing": '_last'
                }
            }
        )

    return s


def filter_results(s, get_params):
    """Фильтрует результат запроса ElasticSearch и выполняет агрегацию для фильтров в сайдбаре."""
    # Возможные фильтры
    filters = [
        {
            'title': 'obj_type',
            'field': 'Document.idObjType'
        },
        {
            'title': 'obj_state',
            'field': 'search_data.obj_state'
        },
        {
            'title': 'registration_status_color',
            'field': 'search_data.registration_status_color'
        },
    ]

    # Если включены платные услуги, то необходимо включить возможность фильтрации по коду MarkCurrentStatusCodeType
    paid_services_settings, created = PaidServicesSettings.objects.get_or_create()
    """ ВРЕМЕННО ОТКРЫТЬ ДОСТУП ВСЕМ """
    # if paid_services_settings.enabled:
    #     filters.append({
    #         'title': 'mark_status',
    #         'field': 'Document.MarkCurrentStatusCodeType.keyword'
    #     })
    # filters.append({
    #     'title': 'mark_status',
    #     'field': 'Document.MarkCurrentStatusCodeType.keyword'
    # })

    # Агрегация без фильтрации
    for item in filters:
        s.aggs.bucket(f"{item['title']}_terms", A('terms', field=item['field'], order={"_key": "asc"}, size=1000))

    aggregations = s.execute().aggregations.to_dict()
    s_ = s

    # Фильтрация
    for item in filters:
        if get_params.get(f"filter_{item['title']}"):
            # Фильтрация в основном запросе
            s = s.filter('terms', **{item['field']: get_params.get(f"filter_{item['title']}")})

        # Применение остальных фильтров в отдельном запросе для последующей агрегации
        s_item = s_
        for item_ in filters:
            if item_ != item and get_params.get(f"filter_{item_['title']}"):
                s_item = s_item.filter(
                    'terms',
                    **{item_['field']: get_params.get(f"filter_{item_['title']}")}
                )

        # Агрегация
        aggregations_item = s_item.execute().aggregations.to_dict()
        item_terms_title = f"{item['title']}_terms"
        for bucket in aggregations[item_terms_title]['buckets']:
            if not list(filter(
                    lambda x: x['key'] == bucket['key'],
                    aggregations_item[item_terms_title]['buckets']
            )):
                aggregations_item[item_terms_title]['buckets'].append(
                    {'key': bucket['key'], 'doc_count': 0}
                )
        aggregations[item_terms_title]['buckets'] = aggregations_item[item_terms_title]['buckets']

    return s, aggregations


def extend_doc_flow(hit):
    """Расширяет секцию DOCFLOW патента документами заявки"""
    # Получение заявки охранного документа
    q = Q(
        'bool',
        must=[
            Q('match', search_data__app_number=hit['search_data']['app_number']),
            Q('match', search_data__obj_state=1)
        ]
    )
    client = Elasticsearch(settings.ELASTIC_HOST, timeout=settings.ELASTIC_TIMEOUT)
    application = Search().using(client).query(q).source(
        excludes=["*.DocBarCode", "*.DOCBARCODE"]
    ).execute()
    if application:
        application = application[0].to_dict()

        try:
            # Объединение стадий
            stages = application['DOCFLOW']['STAGES']
            stages.extend(hit['DOCFLOW']['STAGES'])
            hit['DOCFLOW']['STAGES'] = stages
        except (AttributeError, KeyError):
            pass

        try:
            # Объединение документов
            documents = application['DOCFLOW']['DOCUMENTS']
            documents.extend(hit['DOCFLOW']['DOCUMENTS'])
            hit['DOCFLOW']['DOCUMENTS'] = documents
        except (AttributeError, KeyError):
            pass

        try:
            # Объединение платежей
            payments = application['DOCFLOW']['PAYMENTS']
            payments.extend(hit['DOCFLOW']['PAYMENTS'])
            hit['DOCFLOW']['PAYMENTS'] = payments
        except (AttributeError, KeyError):
            pass

        try:
            # Объединение сборов
            collections = application['DOCFLOW']['COLLECTIONS']
            collections.extend(hit['DOCFLOW']['COLLECTIONS'])
            hit['DOCFLOW']['COLLECTIONS'] = collections
        except (AttributeError, KeyError):
            pass


def get_completed_order(order_id, attempts=10, timeout=2):
    """Ждёт пока обрабатывается заказ. Возвращает False если он не обработался."""
    completed = False
    counter = 0
    while completed is False:
        order = OrderService.objects.get(id=order_id)
        if order.order_completed:
            return order
        else:
            if counter == attempts:
                return False
            counter += 1
            time.sleep(timeout)


def create_selection_inv_um_ld(data_from_json, params, file_path):
    """Формирует документ ворд и сохраняет его на диск"""
    data = dict()
    data['category'] = data_from_json.pop('Category', None)
    data['contracts_comment'] = data_from_json.pop('Contracts_comment', None)
    data['contracts'] = data_from_json.pop('Contracts', [])
    data['collections_comment'] = data_from_json.pop('Collections_comment', None)
    data['collections'] = data_from_json.pop('Collections', [])
    data['documents_comment'] = data_from_json.pop('Documents_comment', None)
    data['documents'] = data_from_json.pop('Documents', [])
    data['publications_comment'] = data_from_json.pop('Publications_comment', None)
    data['publications'] = data_from_json.pop('Publications', [])
    data['signer_comment'] = data_from_json.pop('Signer_comment', None)
    data['signer'] = data_from_json.pop('Signer', None)
    data['INID_51_HTML'] = data_from_json.pop('INID_51_HTML', None)
    data['main_info_comments'] = list(data_from_json.values())[::2]
    data['main_info_values'] = list(data_from_json.values())[1::2]

    # Формирование выписки
    document = Document('selection_templates/template.docx')

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = data['main_info_values'][0]
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.paragraphs[0]
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    run = paragraph.add_run('ВИПИСКА')
    run.bold = True
    run.font.name = 'Times New Roman CYR'
    run.font.size = Pt(18)

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = Pt(18)
    paragraph_format.space_before = 0
    if data['category'] in ('винаходи', 'корисні моделі'):
        run = paragraph.add_run(f"з Державного реєстру патентів України на {data['category']}")
    else:
        run = paragraph.add_run(f"з Державного реєстру свідоцтв України на {data['category']}")
    run.font.name = 'Times New Roman CYR'
    run.font.size = Pt(14)

    table = document.add_table(rows=0, cols=2)
    for i in range(len(data['main_info_comments'])):
        if data['main_info_values'][i] != '':
            row_cells = table.add_row().cells
            row_cells[0].text = data['main_info_comments'][i]
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[0].paragraphs[0].space_after = 0
            row_cells[0].paragraphs[0].space_before = 0
            row_cells[1].text = data['main_info_values'][i]
            row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[1].paragraphs[0].space_after = 0
            row_cells[1].paragraphs[0].space_before = 0

    if params.get('contracts') and data['contracts_comment'] and len(data['contracts']) > 0:
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.space_after = Pt(18)
        paragraph_format.space_before = Pt(18)
        run = paragraph.add_run(data['contracts_comment'])
        run.bold = True
        run.font.name = 'Times New Roman CYR'
        run.font.size = Pt(12)

        for item in data['contracts']:
            item_values = list(item.values())
            comments = item_values[::2]
            values = item_values[1::2]
            table = document.add_table(rows=0, cols=2)
            for i in range(len(comments)):
                if comments[i] != '' and values[i] != '':
                    row_cells = table.add_row().cells
                    row_cells[0].text = comments[i]
                    row_cells[0].paragraphs[0].runs[0].font.bold = True
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[0].paragraphs[0].space_after = 0
                    row_cells[0].paragraphs[0].space_before = 0
                    row_cells[1].text = values[i]
                    row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[1].paragraphs[0].space_after = 0
                    row_cells[1].paragraphs[0].space_before = 0
                    if i == len(comments) - 1:
                        set_cell_border(
                            row_cells[1],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
                        set_cell_border(
                            row_cells[0],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
            document.add_paragraph()

    if params.get('collections') and data['collections_comment'] and len(data['collections']) > 0:
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(data['collections_comment'])
        run.bold = True
        run.font.name = 'Times New Roman CYR'
        run.font.size = Pt(12)

        for collection in data['collections']:
            res = dict()
            for key, value in collection.items():
                if key != 'Summs':
                    res[key] = value
                else:
                    res_summs = dict()
                    summs = value[0]
                    for key_summs, value_summs in summs.items():
                        if key_summs != 'Payments':
                            res_summs[key_summs] = value_summs
                        else:
                            if len(value_summs):
                                payments = value_summs[0]
                                payments.pop('nope', None)
                                res_summs.update(payments)
                    res.update(res_summs)

            item_values = list(res.values())
            comments = item_values[::2]
            values = item_values[1::2]
            table = document.add_table(rows=0, cols=2)
            for i in range(len(comments)):
                if comments[i] != '' and values[i] != '':
                    row = table.add_row()
                    row_cells = row.cells
                    row_cells[0].text = comments[i]
                    row_cells[0].paragraphs[0].runs[0].font.bold = True
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[0].paragraphs[0].space_after = 0
                    row_cells[0].paragraphs[0].space_before = 0
                    row_cells[1].text = values[i]
                    row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[1].paragraphs[0].space_after = 0
                    row_cells[1].paragraphs[0].space_before = 0
                    if i == len(comments) - 1:
                        set_cell_border(
                            row_cells[1],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
                        set_cell_border(
                            row_cells[0],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
            document.add_paragraph()

    if params.get('letters') and data['documents_comment'] and len(data['documents']) > 0:
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(data['documents_comment'])
        run.bold = True
        run.font.name = 'Times New Roman CYR'
        run.font.size = Pt(12)

        for item in data['documents']:
            item_values = list(item.values())
            comments = item_values[::2]
            values = item_values[1::2]
            table = document.add_table(rows=0, cols=2)
            for i in range(len(comments)):
                if comments[i] != '' and values[i] != '':
                    row_cells = table.add_row().cells
                    row_cells[0].text = comments[i]
                    row_cells[0].paragraphs[0].runs[0].font.bold = True
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[0].paragraphs[0].space_after = 0
                    row_cells[0].paragraphs[0].space_before = 0
                    row_cells[1].text = values[i]
                    row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[1].paragraphs[0].space_after = 0
                    row_cells[1].paragraphs[0].space_before = 0
                    if i == len(comments) - 1:
                        set_cell_border(
                            row_cells[1],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
                        set_cell_border(
                            row_cells[0],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
            document.add_paragraph()

    if params.get('another') and data['publications_comment'] and len(data['publications']) > 0:
        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(data['publications_comment'])
        run.bold = True
        run.font.name = 'Times New Roman CYR'
        run.font.size = Pt(12)

        for item in data['publications']:
            item_values = list(item.values())
            comments = item_values[::2]
            values = item_values[1::2]
            table = document.add_table(rows=0, cols=2)
            for i in range(len(comments)):
                if comments[i] != '' and values[i] != '':
                    row_cells = table.add_row().cells
                    row_cells[0].text = comments[i]
                    row_cells[0].paragraphs[0].runs[0].font.bold = True
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[0].paragraphs[0].space_after = 0
                    row_cells[0].paragraphs[0].space_before = 0
                    row_cells[1].text = values[i]
                    row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
                    row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
                    row_cells[1].paragraphs[0].space_after = 0
                    row_cells[1].paragraphs[0].space_before = 0
                    if i == len(comments) - 1:
                        set_cell_border(
                            row_cells[1],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
                        set_cell_border(
                            row_cells[0],
                            bottom={"sz": 12, "color": "#000000", "val": "single"},
                        )
            document.add_paragraph()

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(18)
    paragraph_format.space_before = Pt(40)
    now = datetime.datetime.now()
    now = now.strftime("%d.%m.%Y")
    run = paragraph.add_run(f"Виписка видана станом на {now} р.")
    run.font.name = 'Times New Roman CYR'
    run.bold = True
    run.font.size = Pt(12)

    table = document.add_table(rows=0, cols=2)
    row_cells = table.add_row().cells
    row_cells[0].text = data['signer']['SignerPost']
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row_cells[0].paragraphs[0].space_after = 0
    row_cells[0].paragraphs[0].space_before = 0
    row_cells[1].text = data['signer']['SignerName']
    row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman CYR'
    row_cells[1].paragraphs[0].runs[0].font.bold = True
    row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    row_cells[1].paragraphs[0].space_after = 0
    row_cells[1].paragraphs[0].space_before = 0

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(18)
    run = paragraph.add_run('М.П.')
    run.font.name = 'Times New Roman CYR'
    run.bold = True
    run.font.size = Pt(12)

    # Сохранение в файл
    document.save(file_path)


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def get_data_for_selection_tm(hit):
    """Возвращает данные для выписки по знаку для товаров и услуг"""
    data = dict()
    data['biblio'] = dict()
    # (111) Номер реєстрації знака, який є номером свідоцтва
    data['biblio']['i_111'] = {
        'key': '111',
        'value': hit.search_data.protective_doc_number,
    }
    # (151) Дата реєстрації знака
    data['biblio']['i_151'] = {
        'key': '151',
        'value': datetime.datetime.strptime(hit.search_data.rights_date, '%Y-%m-%d').strftime('%d.%m.%Y')
    }
    # (181) Очікувана дата закінчення строку дії свідоцтва
    data['biblio']['i_181'] = {
        'key': '181',
        'value': datetime.datetime.strptime(hit.TradeMark.TrademarkDetails.ExpiryDate, '%Y-%m-%d').strftime('%d.%m.%Y')
    }
    # (210) Номер заявки
    data['biblio']['i_210'] = {
        'key': '210',
        'value': hit.search_data.app_number
    }
    # (220) Дата подання заявки
    data['biblio']['i_220'] = {
        'key': '220',
        'value': datetime.datetime.strptime(hit.search_data.app_date, '%Y-%m-%d').strftime('%d.%m.%Y')
    }
    # (450) Бюлетень
    data['biblio']['i_450'] = {
        'key': '450',
        'value': f"{hit.TradeMark.TrademarkDetails.PublicationDetails.Publication.PublicationDate}. "
                 f"Бюл. {hit.TradeMark.TrademarkDetails.PublicationDetails.Publication.PublicationIdentifier}"
    }
    # (591) зазначення кольору чи поєднання кольорів, які охороняються
    data['biblio']['i_591'] = {
        'key': '591',
        'value': []
    }
    try:
        for color in hit.TradeMark.TrademarkDetails.MarkImageDetails.MarkImage.MarkImageColourClaimedText:
            data['biblio']['i_591']['value'].append(color['#text'])
    except AttributeError:
        pass

    # (732) Ім'я або повне найменування та адреса власника (власників) свідоцтва
    data['biblio']['i_732'] = {
        'key': '732',
        'value': []
    }
    for holder in hit.TradeMark.TrademarkDetails.HolderDetails.Holder:
        data['biblio']['i_732']['value'].append(
            {
                'name': holder.HolderAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine,
                'address': holder.HolderAddressBook.FormattedNameAddress.Address.FreeFormatAddress.FreeFormatAddressLine
            }
        )
    # (511) Індекс (індекси) Міжнародної класифікації товарів і послуг для реєстрації знаків та перелік товарів і послуг
    data['biblio']['i_511'] = {
        'key': '511',
        'value': []
    }
    for cls in hit.TradeMark.TrademarkDetails.GoodsServicesDetails.GoodsServices.ClassDescriptionDetails.ClassDescription:
        data['biblio']['i_511']['value'].append(
            {
                'cls': cls.ClassNumber,
                'terms': '; '.join([term.ClassificationTermText for term in cls.ClassificationTermDetails.ClassificationTerm])
            }
        )
    # (540) Зображення знака
    splitted_path = hit.Document.filesPath.replace("\\", "/").split('/')
    splitted_path_len = len(splitted_path)
    data['biblio']['i_540'] = {
        'key': '540',
        'value': f"{settings.MEDIA_ROOT}"
                 f"{splitted_path[splitted_path_len-4]}/"
                 f"{splitted_path[splitted_path_len-3]}/"
                 f"{splitted_path[splitted_path_len-2]}/"
                 f"{hit.TradeMark.TrademarkDetails.MarkImageDetails.MarkImage.MarkImageFilename}"
    }

    data['transactions'] = []
    try:
        data['transactions'] = hit.TradeMark.Transactions.Transaction
    except AttributeError:
        pass

    data['signer'] = {
        'post': 'Заступник Міністра розвитку економіки, торгівлі та сільського господарства України',
        'name': 'Д.О.Романович'
    }

    return data


def create_selection_tm(data, params, file_path):
    """Формирует документ ворд и сохраняет его на диск."""
    # Формирование выписки
    document = Document('selection_templates/template.docx')

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = data['biblio']['i_111']['value']
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = document.paragraphs[0]
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    run = paragraph.add_run('Виписка')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    run = paragraph.add_run("з Державного реєстру свідоцтв України")
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    run = paragraph.add_run("на знаки для товарів та послуг відносно свідоцтва № ")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run = paragraph.add_run(data['biblio']['i_111']['value'])
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_after = Pt(18)
    paragraph_format.space_before = 0
    now = datetime.datetime.now()
    now = now.strftime("%d.%m.%Y")
    run = paragraph.add_run(f"станом на {now} р.")
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    # Библиография
    table = document.add_table(rows=0, cols=2)
    row_cells = table.add_row().cells
    table.cell(0, 0).width = Cm(20)
    for key, value in data['biblio'].items():
        if value['key'] not in ('540', '511') and value['value']:
            if value['key'] == '111':
                paragraph = row_cells[0].paragraphs[0]
            else:
                paragraph = row_cells[0].add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = 0
            paragraph_format.space_before = 0
            paragraph_format = paragraph.paragraph_format
            paragraph_format.first_line_indent = Cm(-1.5)
            paragraph_format.left_indent = Cm(1.5)
            run = paragraph.add_run(f"({value['key']})\t")
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run.bold = True
            if value['key'] == '591':
                run = paragraph.add_run(', '.join(value['value']))
            elif value['key'] == '732':
                for i in range(len(value['value'])):
                    if i > 0:
                        paragraph = row_cells[0].add_paragraph()
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Cm(1.5)
                        paragraph_format.space_after = 0
                        paragraph_format.space_before = 0
                    run = paragraph.add_run(f"{value['value'][i]['name']};")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                    paragraph = row_cells[0].add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.left_indent = Cm(1.5)
                    paragraph_format.space_after = 0
                    paragraph_format.space_before = 0
                    run = paragraph.add_run(f"{value['value'][i]['address']}")
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
                    run.bold = True
            else:
                run = paragraph.add_run(f"{value['value']}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if value['key'] in ('111', '151'):
                run.bold = True

    paragraph = row_cells[1].paragraphs[0]
    run = paragraph.add_run('(540)')
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Cm(1.5)
    paragraph_format.space_before = 0
    paragraph = row_cells[1].add_paragraph()
    run = paragraph.add_run()
    run.add_picture(data['biblio']['i_540']['value'], width=Cm(5))

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = 0
    paragraph_format.space_before = 0
    paragraph_format.first_line_indent = Cm(-1.5)
    paragraph_format.left_indent = Cm(1.5)
    run = paragraph.add_run("(511)")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True

    for i_511 in data['biblio']['i_511']['value']:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = 0
        paragraph_format.space_before = 0
        paragraph_format.first_line_indent = Cm(-1.5)
        paragraph_format.left_indent = Cm(1.5)
        run = paragraph.add_run(f"Кл. {i_511['cls']}:\t")
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run = paragraph.add_run(i_511['terms'])
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # Сповіщення
    if params.get('transactions') and len(data['transactions']) > 0:
        for transaction in data['transactions']:
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Cm(0.5)
            paragraph_format.left_indent = Cm(1.5)
            run = paragraph.add_run(transaction['@name'])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.underline = True

            # тип лицензии в заголовке оповещения
            try:
                license_kind = 'Невиключна' if transaction.TransactionBody.LicenceKind == 'Nonexclusive' else 'Виключна'
                run = paragraph.add_run(f" ({license_kind} ліцензія)")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.underline = True
            except AttributeError:
                pass

            # "рішення №" в заголовке оповещения
            try:
                decision_number = transaction.TransactionBody.DecisionDetails.Decision.Number
                decision_number = decision_number.replace('№ ', '').replace('№', '')
                run = paragraph.add_run(f", рішення № {decision_number}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.underline = True
            except AttributeError:
                pass

            # "від " в заголовке оповещения
            try:
                decision_date = transaction.TransactionBody.DecisionDetails.Decision.Date
                decision_date = datetime.datetime.strptime(decision_date, '%Y-%m-%d').strftime('%d.%m.%Y')
                run = paragraph.add_run(f" від {decision_date}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.underline = True
            except AttributeError:
                pass

            # 141
            try:
                termination_date = transaction.TransactionBody.TerminationDate
                termination_date = datetime.datetime.strptime(termination_date, '%Y-%m-%d').strftime('%d.%m.%Y')
                paragraph = document.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = 0
                paragraph_format.space_after = 0
                paragraph_format.left_indent = Cm(1.5)
                paragraph_format.first_line_indent = Cm(-1.5)
                run = paragraph.add_run('(141)')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run = paragraph.add_run(f"\t{termination_date}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            except AttributeError:
                pass

            # 186
            prolongation_expiry_date = None
            try:
                prolongation_expiry_date = transaction.TransactionBody.ProlongationExpiryDate
                prolongation_expiry_date = datetime.datetime.strptime(prolongation_expiry_date, '%Y-%m-%d').strftime('%d.%m.%Y')
                paragraph = document.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = 0
                paragraph_format.space_after = 0
                paragraph_format.left_indent = Cm(1.5)
                paragraph_format.first_line_indent = Cm(-1.5)
                run = paragraph.add_run('(186)')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run = paragraph.add_run(f"\t{prolongation_expiry_date}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
            except AttributeError:
                pass

            # 450
            try:
                bulletin_date = transaction['@bulletinDate']
                bulletin_date = datetime.datetime.strptime(bulletin_date, '%Y-%m-%d').strftime('%d.%m.%Y')
                bulletin_number = transaction['@bulletinNumber']
                paragraph = document.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = 0
                paragraph_format.space_after = 0
                paragraph_format.left_indent = Cm(1.5)
                paragraph_format.first_line_indent = Cm(-1.5)
                run = paragraph.add_run('(450)')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run = paragraph.add_run(f"\t{bulletin_date}. Бюл. № {bulletin_number}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            except AttributeError:
                pass

            # 580
            if not prolongation_expiry_date:
                try:
                    bulletin_date = transaction['@bulletinDate']
                    bulletin_date = datetime.datetime.strptime(bulletin_date, '%Y-%m-%d').strftime('%d.%m.%Y')
                    paragraph = document.add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = 0
                    paragraph_format.space_after = 0
                    paragraph_format.left_indent = Cm(1.5)
                    paragraph_format.first_line_indent = Cm(-1.5)
                    run = paragraph.add_run('(580)')
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                    run = paragraph.add_run(f"\t{bulletin_date}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                except AttributeError:
                    pass

            # 732
            holder = None
            try:
                holder = transaction.TransactionBody.HolderDetails.NewHolder
            except AttributeError:
                try:
                    holder = transaction.TransactionBody.HolderDetails.Holder
                except AttributeError:
                    pass
            finally:
                if holder:
                    paragraph = document.add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = 0
                    paragraph_format.space_after = 0
                    paragraph_format.left_indent = Cm(1.5)
                    paragraph_format.first_line_indent = Cm(-1.5)
                    run = paragraph.add_run('(732)\t')
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.bold = True
                    for i in range(len(holder)):
                        if i > 0:
                            paragraph = document.add_paragraph()
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.left_indent = Cm(1.5)
                            paragraph_format.space_after = 0
                            paragraph_format.space_before = 0
                        run = paragraph.add_run(f"{holder[i].ApplicantAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine};")
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = True
                        paragraph = document.add_paragraph()
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.left_indent = Cm(1.5)
                        paragraph_format.space_after = 0
                        paragraph_format.space_before = 0
                        run = paragraph.add_run(f"{holder[i].ApplicantAddressBook.FormattedNameAddress.Address.FreeFormatAddressLine} "
                                                f"({holder[i].ApplicantAddressBook.FormattedNameAddress.Address.AddressCountryCode})")
                        run.font.size = Pt(12)
                        run.font.name = 'Times New Roman'
                        run.bold = True

    # Подпись
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(12)
    table = document.add_table(rows=0, cols=2)
    row_cells = table.add_row().cells
    row_cells[0].text = data['signer']['post']
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    row_cells[0].paragraphs[0].space_after = 0
    row_cells[0].paragraphs[0].space_before = 0
    row_cells[1].text = data['signer']['name']
    row_cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
    row_cells[1].paragraphs[0].runs[0].font.bold = True
    row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    row_cells[1].paragraphs[0].space_after = 0
    row_cells[1].paragraphs[0].space_before = 0

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(18)
    run = paragraph.add_run('М.П.')
    run.font.name = 'Times New Roman'
    run.bold = True
    run.font.size = Pt(12)

    document.save(file_path)


def user_has_access_to_docs_decorator(f):
    """Декоратор для проверки у пользователя наличия прав на скачивание документа(ов)."""
    def wrapper(*args, **kw):
        id_app_number = args[0].POST.get('id_app_number') or kw.get('id_app_number')
        if user_has_access_to_docs(args[0].user, id_app_number):
            return f(*args, **kw)
        else:
            return HttpResponse('401 Unauthorized', status=401)
    return wrapper


def user_has_access_to_docs(user, hit):
    """Возвращает признак доступности документа(ов)"""
    # Проверка на принадлженость пользователя к роли суперадмина или к ВИП-роли
    if not user.is_anonymous:
        if user.is_vip():
            return True
        else:
            allowed_persons = []

            for person_type in ('applicant', 'inventor', 'owner', 'agent'):
                try:
                    if hit['search_data'].get(person_type):
                        allowed_persons += [x['name'] for x in hit['search_data'][person_type]]
                except KeyError:
                    pass

            # Адреса для листування (ТМ)
            try:
                allowed_persons.append(
                    hit['TradeMark']['TrademarkDetails']['CorrespondenceAddress']['CorrespondenceAddressBook'][
                        'Name']['FreeFormatNameLine']
                )
            except KeyError:
                pass

            # Адреса для листування (ПЗ)
            try:
                allowed_persons.append(
                    hit['Design']['DesignDetails']['CorrespondenceAddress']['CorrespondenceAddressBook'][
                        'FormattedNameAddress']['Name']['FreeFormatName']['FreeFormatNameDetails']['FreeFormatNameLine']
                )
            except KeyError:
                pass

            # Адреса для листування (заявки на винаходи, корисні моделі)
            try:
                allowed_persons.append(hit['Claim']['I_98'])
            except KeyError:
                pass

            # Адреса для листування (охоронні документи на винаходи, корисні моделі)
            try:
                allowed_persons.append(hit['Patent']['I_98'])
            except KeyError:
                pass

            # Имена пользователя (если роль - патентный поверенный, то необходимо проверять несколько имён)
            user_names = []
            if hasattr(user, 'certificateowner'):
                user_names = [user.certificateowner.pszSubjFullName.strip()]
            elif user.is_patent_attorney():
                user_names.append(f"{user.last_name} {user.first_name}".strip())
                for patent_attroney in user.patentattorney_set.all():
                    user_names.append(patent_attroney.name)

            # Проверка на вхождение
            return any([person for person in allowed_persons for user_name in user_names if
                        user_name.upper() in person.upper()])

    return False


def user_has_access_to_tm_app(user, hit):
    """Возвращает признак доступности заявки на знак для товаров и услуг пользователю."""
    if not type(hit) is dict:
        item = hit.to_dict()
        item['meta'] = hit.meta.to_dict()
        hit = item
    try:
        return hit['Document'].get('MarkCurrentStatusCodeType') != '1000' \
               or user_has_access_to_docs(user, hit) \
               or user in IpcAppList.objects.get(id=hit['meta']['id']).users_with_access.all()
    except IpcAppList.DoesNotExist:
        return False


def create_search_res_doc(data, file_path):
    """Формировние Excel-файла с результатами поиска."""
    workbook = xlsxwriter.Workbook(file_path)
    worksheet = workbook.add_worksheet()

    # Заголовки
    titles = [
        _("Тип об'єкта промислової власності"),
        _("Стан об'єкта промислової власності"),
        _("Номер заявки"),
        _("Дата подання заявки"),
        _("Номер охоронного документа"),
        _("Дата охоронного документа"),
        _("Ключові слова"),
        _("Заявник"),
        _("Власник"),
        _('Винахідник\автор'),
        _("Представник"),
        _("МПК"),
        _("МКТП"),
        _("МКПЗ"),
        _("(441) Дата публікації відомостей про заявку та номер бюлетня"),
        _("Зображення ТМ"),
    ]

    bold = workbook.add_format({'bold': True})
    worksheet.write_row(0, 0, titles, bold)

    for row_num, row_data in enumerate(data):
        for col_num, col_data in enumerate(row_data):
            if col_num == 15 and col_data and os.path.exists(col_data):
                worksheet.insert_image(
                    row_num + 1,
                    col_num,
                    col_data,
                    {'x_offset': 2, 'y_offset': 2, 'x_scale': 0.3, 'y_scale': 0.3}
                )
            else:
                worksheet.write(row_num + 1, col_num, col_data)

    workbook.close()


def prepare_data_for_search_report(s, lang_code, user=None):
    """Подготавливает данные для файла Excel."""
    obj_states = [_('Заявка'), _('Охоронний документ')]
    obj_types = ObjType.objects.order_by('id').values_list('id', f"obj_type_{lang_code}")
    data = list()
    for h in s.params(size=1000, preserve_order=True).scan():
        obj_type = next(filter(lambda item: item[0] == h.Document.idObjType, obj_types), None)[1]
        obj_state = obj_states[h.search_data.obj_state - 1]

        if is_app_limited(h.to_dict(), user):
            # Если библиографические данные заявки не публикуются
            data.append([
                obj_type,
                obj_state,
                h.search_data.app_number,
                '',
                '',
                '',
                '',
                '',
                '',
                '',
                '',
                '',
                '',
            ])
        else:
            app_date = datetime.datetime.strptime(h.search_data.app_date[:10], '%Y-%m-%d').strftime('%d.%m.%Y') \
                if hasattr(h.search_data, 'app_date') else ''
            rights_date = datetime.datetime.strptime(h.search_data.rights_date, '%Y-%m-%d').strftime(
                '%d.%m.%Y') if h.search_data.rights_date else ''
            title = ';\r\n'.join(h.search_data.title) if iterable(h.search_data.title) else h.search_data.title
            applicant = get_app_applicant(h)
            owner = get_app_owner(h)
            inventor = get_app_inventor(h)
            if hasattr(h.search_data, 'agent') and h.search_data.agent:
                agent = ';\r\n'.join([x.name for x in h.search_data.agent])
            else:
                agent = ''
            ipc_indexes = get_app_ipc_indexes(h)
            nice_indexes = get_app_nice_indexes(h)
            icid = get_app_icid(h)
            if h.Document.idObjType in (4, 9, 14):
                code_441 = get_441_code(h)
                image = get_tm_image_path(h)
            else:
                code_441 = ''
                image = ''

            data.append([
                    obj_type,
                    obj_state,
                    h.search_data.app_number,
                    app_date,
                    h.search_data.protective_doc_number,
                    rights_date,
                    title,
                    applicant,
                    owner,
                    inventor,
                    agent,
                    ipc_indexes,
                    nice_indexes,
                    icid,
                    code_441,
                    image,
                ])
    return data


def get_tm_image_path(app):
    """Возвращает путь к изображению ТМ на диске."""
    if type(app) is not dict:
        app = app.to_dict()
    splitted_path = app['Document']['filesPath'].replace("\\", "/").split('/')
    splitted_path_len = len(splitted_path)

    if app['Document']['idObjType'] == 4:
        image_name = app['TradeMark']['TrademarkDetails']['MarkImageDetails']['MarkImage']['MarkImageFilename']
        return f"{settings.MEDIA_ROOT}/" \
               f"{splitted_path[splitted_path_len-4]}" \
               f"/{splitted_path[splitted_path_len-3]}/" \
               f"{splitted_path[splitted_path_len-2]}/{image_name}"
    else:
        # ТМ Мадрид
        image_name = f"{app['search_data']['protective_doc_number']}.jpg"
        return f"{settings.MEDIA_ROOT}/" \
               f"{splitted_path[splitted_path_len-5].upper()}/" \
               f"{splitted_path[splitted_path_len-4]}/" \
               f"{splitted_path[splitted_path_len-3]}/" \
               f"{splitted_path[splitted_path_len-2]}/{image_name}"


def get_441_code(app):
    """Возвращает 441 код ТМ."""
    if type(app) is not dict:
        app = app.to_dict()

    if app['Document']['idObjType'] in (4, 9, 14):
        try:
            if app['Document']['idObjType'] == 4:
                code_441 = app['TradeMark']['TrademarkDetails']['Code_441']
            else:
                code_441 = app['MadridTradeMark']['TradeMarkDetails']['Code_441']
            code_441_formatted = datetime.datetime.strptime(code_441, '%Y-%m-%d').strftime('%d.%m.%Y')
        except KeyError:
            return ''

        try:
            obj = ClListOfficialBulletinsIp.objects.get(date_from__lte=code_441, date_to__gte=code_441)
        except ClListOfficialBulletinsIp.DoesNotExist:
            return code_441_formatted
        else:
            return f"{code_441_formatted}, бюл. №{obj.bul_number}"
    else:
        return ''


def get_app_applicant(app):
    """Возвращает строку с заявителями (с кодами стран)"""
    # Изобретения, полезные модели, топографии
    applicants = []

    if app.Document.idObjType in (1, 2, 3):
        if app.search_data.obj_state == 1:
            biblio_data = app.Claim.to_dict()
        else:
            biblio_data = app.Patent.to_dict()

        if biblio_data.get('I_71'):
            i_71 = [i for n, i in enumerate(biblio_data['I_71']) if i not in biblio_data['I_71'][n + 1:]]
            for item in i_71:
                item.pop('EDRPOU', '')
                item_values = sorted(list(item.values()), key=len)
                applicants.append(f"{item_values[1]} [{item_values[0]}]")
    # Знаки для товаров и услуг
    elif app.Document.idObjType == 4:
        try:
            for item in app.TradeMark.TrademarkDetails.ApplicantDetails.Applicant:
                applicants.append(
                    f"{item.ApplicantAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.ApplicantAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    # Пром. образцы
    elif app.Document.idObjType == 6:
        try:
            for item in app.Design.DesignDetails.ApplicantDetails.Applicant:
                applicants.append(
                    f"{item.ApplicantAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.ApplicantAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    # Авторское право
    elif app.Document.idObjType in (10, 13):
        try:
            for item in app.Certificate.CopyrightDetails.ApplicantDetails.Applicant:
                applicants.append(
                    f"{item.ApplicantAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.ApplicantAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    return ';\r\n'.join(applicants)


def get_app_owner(app):
    """Возвращает строку с владельцами (с кодами стран)"""
    # Изобретения, полезные модели, топографии
    owners = []

    if app.Document.idObjType in (1, 2, 3):
        if app.search_data.obj_state == 1:
            biblio_data = app.Claim.to_dict()
        else:
            biblio_data = app.Patent.to_dict()

        if biblio_data.get('I_73'):
            i_73 = [i for n, i in enumerate(biblio_data['I_73']) if i not in biblio_data['I_73'][n + 1:]]
            for item in i_73:
                item.pop('EDRPOU', '')
                item_values = sorted(list(item.values()), key=len)
                owners.append(f"{item_values[1]} [{item_values[0]}]")

    # Знаки для товаров и услуг
    elif app.Document.idObjType == 4:
        try:
            for item in app.TradeMark.TrademarkDetails.HolderDetails.Holder:
                owners.append(
                    f"{item.HolderAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.HolderAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    # КЗПТ
    elif app.Document.idObjType == 5:
        try:
            for item in app.Geo.GeoDetails.HolderDetails.Holder:
                owners.append(
                    f"{item.HolderAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.HolderAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    # Пром. образцы
    elif app.Document.idObjType == 6:
        try:
            for item in app.Design.DesignDetails.HolderDetails.Holder:
                owners.append(
                    f"{item.HolderAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.HolderAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    # ТМ (Мадрид)
    elif app.Document.idObjType in (9, 14):
        owner = app.MadridTradeMark.TradeMarkDetails.HOLGR.NAME.NAMEL
        for item in app.MadridTradeMark.TradeMarkDetails.HOLGR.ADDRESS.ADDRL:
            owner += f" {item}"
        if app.MadridTradeMark.TradeMarkDetails.HOLGR.ADDRESS.COUNTRY:
            owner += f"[{app.MadridTradeMark.TradeMarkDetails.HOLGR.ADDRESS.COUNTRY}]"
        owners.append(owner)

    # Авторское право
    elif app.Document.idObjType in (10, 13):
        try:
            for item in app.Certificate.CopyrightDetails.HolderDetails.Holder:
                owners.append(
                    f"{item.HolderAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.HolderAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    return ';\r\n'.join(owners)


def get_app_inventor(app):
    """Возвращает строку с изобреттелями (с кодами стран)"""
    # Изобретения, полезные модели, топографии
    inventors = []

    if app.Document.idObjType in (1, 2, 3):
        if app.search_data.obj_state == 1:
            biblio_data = app.Claim.to_dict()
        else:
            biblio_data = app.Patent.to_dict()

        if biblio_data.get('I_72'):
            i_72 = [i for n, i in enumerate(biblio_data['I_72']) if i not in biblio_data['I_72'][n + 1:]]
            for item in i_72:
                item.pop('EDRPOU', '')
                item_values = sorted(list(item.values()), key=len)
                inventors.append(f"{item_values[1]} [{item_values[0]}]")

    # Пром. образцы
    elif app.Document.idObjType == 6:
        try:
            for item in app.Design.DesignDetails.DesignerDetails.Designer:
                inventors.append(
                    f"{item.DesignerAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.DesignerAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    # Авторское право
    elif app.Document.idObjType in (10, 13):
        try:
            for item in app.Certificate.CopyrightDetails.AuthorDetails.Author:
                inventors.append(
                    f"{item.AuthorAddressBook.FormattedNameAddress.Name.FreeFormatName.FreeFormatNameDetails.FreeFormatNameLine} "
                    f"[{item.AuthorAddressBook.FormattedNameAddress.Address.AddressCountryCode}]"
                )
        except AttributeError:
            pass

    return ';\r\n'.join(inventors)


def get_app_ipc_indexes(app):
    """Возвращает строку с индексами МПК заявки"""
    if app.Document.idObjType in (1, 2):
        ipc_indexes = []
        if app.search_data.obj_state == 1:
            biblio_data = app.Claim
        else:
            biblio_data = app.Patent
        try:
            for IPC in biblio_data.IPC:
                ipc_indexes.append(IPC)
        except AttributeError:
            pass
        else:
            return '; '.join(ipc_indexes)
    return ''


def get_app_nice_indexes(app):
    """Возвращает строку с индексами МКТП заявки"""
    if app.Document.idObjType == 4 and hasattr(app.TradeMark.TrademarkDetails, 'GoodsServicesDetails'):
        nice_indexes = []
        for cls in app.TradeMark.TrademarkDetails.GoodsServicesDetails.GoodsServices.ClassDescriptionDetails.ClassDescription:
            nice_indexes.append(str(cls.ClassNumber))
        return '; '.join(nice_indexes)

    elif app.Document.idObjType in (9, 14):
        nice_indexes = []
        for cls in app.MadridTradeMark.TradeMarkDetails.BASICGS.GSGR:
            nice_indexes.append(str(cls['@NICCLAI']))
        return '; '.join(nice_indexes)

    return ''


def get_app_icid(app):
    """Возвращает строку с индексами МКПЗ заявки"""
    if app.Document.idObjType == 6:
        icid = []
        try:
            for i in app.Design.DesignDetails.IndicationProductDetails:
                icid.append(str(i.Class))
        except AttributeError:
            pass
        else:
            return '; '.join(icid)
    return ''


def get_transactions_types(id_obj_type):
    """Возвращает возможные типы транзакций для определённого типа объекта."""
    client = Elasticsearch(settings.ELASTIC_HOST, timeout=settings.ELASTIC_TIMEOUT)
    q = Q(
        'bool',
        must=[Q('match', Document__idObjType=id_obj_type)],
    )
    if id_obj_type in (1, 2, 3):
        # Изобретения, полезные модели, топографии
        field = 'TRANSACTIONS.TRANSACTION.PUBLICATIONNAME.keyword'
        nested = 'TRANSACTIONS.TRANSACTION'
    elif id_obj_type == 4:
        # Знаки для товаров и услуг
        field = 'TradeMark.Transactions.Transaction.@name.keyword'
        nested = 'TradeMark.Transactions.Transaction'
    elif id_obj_type == 5:
        # КЗПТ
        field = 'Geo.Transactions.Transaction.@name.keyword'
        nested = 'Geo.Transactions.Transaction'
    else:
        # Пром. образцы
        field = 'Design.Transactions.Transaction.@name.keyword'
        nested = 'Design.Transactions.Transaction'

    s = Search().using(client).query(q).extra(size=0)
    s.aggs.bucket('transactions_types', Nested(path=nested)).bucket('transactions_types', Terms(field=field, size=1000, order={"_key": "asc"}))

    try:
        return [x['key'] for x in s.execute().aggregations.to_dict()['transactions_types']['transactions_types']['buckets']]
    except KeyError:
        return []


def get_search_in_transactions(search_params):
    client = Elasticsearch(settings.ELASTIC_HOST, timeout=settings.ELASTIC_TIMEOUT)

    transaction_path = ''
    transaction_type_field = ''
    transaction_date_field = ''
    if search_params['obj_type'] in (1, 2, 3):
        transaction_path = 'TRANSACTIONS.TRANSACTION'
        transaction_type_field = 'TRANSACTIONS.TRANSACTION.PUBLICATIONNAME.keyword'
        transaction_date_field = 'TRANSACTIONS.TRANSACTION.BULLETIN_DATE'
    elif search_params['obj_type'] == 4:
        transaction_path = 'TradeMark.Transactions.Transaction'
        transaction_type_field = 'TradeMark.Transactions.Transaction.@name.keyword'
        transaction_date_field = 'TradeMark.Transactions.Transaction.@bulletinDate'
    elif search_params['obj_type'] == 5:
        transaction_path = 'Geo.Transactions.Transaction'
        transaction_type_field = 'Geo.Transactions.Transaction.@name.keyword'
        transaction_date_field = 'Geo.Transactions.Transaction.@bulletinDate'
    elif search_params['obj_type'] == 6:
        transaction_path = 'Design.Transactions.Transaction'
        transaction_type_field = 'Design.Transactions.Transaction.@name.keyword'
        transaction_date_field = 'Design.Transactions.Transaction.@bulletinDate'

    if transaction_type_field and transaction_date_field:
        qs = Q(
            'bool',
            must=[
                Q('match', search_data__obj_state=2),
                Q('match', Document__idObjType=search_params['obj_type']),
            ]
        )
        qs &= Q(
            'nested',
            path=transaction_path,
            query=Q(
                'bool',
                must=[
                    Q(
                        'query_string',
                        query=' OR '.join([f"\"{x}\"" for x in search_params['transaction_type']]),
                        default_field=transaction_type_field,
                        default_operator='AND'
                    ),
                    Q('range', **{transaction_date_field: {
                        'gte': search_params['date']['date_from'],
                        'lte': search_params['date']['date_to']}
                    })
                ]
            )
        )

        qs = filter_bad_apps(qs)

        s = Search(using=client, index=settings.ELASTIC_INDEX_NAME).source(
            excludes=["*.DocBarCode", "*.DOCBARCODE"]
        ).query(qs)

        return s

    return False


def sort_doc_flow(hit):
    """Сортировка документов по дате."""
    # Изобретения, полезные модели, топографии
    if hit['Document']['idObjType'] in (1, 2, 3):
        if hit.get('DOCFLOW'):
            hit['DOCFLOW'].get('DOCUMENTS', []).sort(
                key=lambda document: time.strptime(document['DOCRECORD'].get(
                    'DOCREGDATE', document['DOCRECORD'].get(
                        'DOCSENDINGDATE', '1970-01-01'
                    )
                ), "%Y-%m-%d")
            )
            hit['DOCFLOW'].get('PAYMENTS', []).sort(
                key=lambda document: time.strptime(document['PFRECORD'].get(
                    'PFDATE', '1970-01-01'
                ), "%Y-%m-%d")
            )
            hit['DOCFLOW'].get('COLLECTIONS', []).sort(
                key=lambda document: time.strptime(document['CLRECORD'].get(
                    'CLDATEBEGIN', '1970-01-01'
                ), "%Y-%m-%d")
            )

    # Знаки для товаров и услуг
    elif hit['Document']['idObjType'] == 4:
        if hit['TradeMark'].get('DocFlow'):
            hit['TradeMark']['DocFlow']['Documents'].sort(
                key=lambda document: time.strptime(document['DocRecord'].get(
                    'DocRegDate', '1970-01-01'
                ), "%Y-%m-%d")
            )
        if hit['TradeMark'].get('PaymentDetails'):
            hit['TradeMark']['PaymentDetails']['Payment'].sort(
                key=lambda document: time.strptime(document.get(
                    'PaymentDate', '1970-01-01'
                ), "%Y-%m-%d")
            )

    # Пром образцы
    elif hit['Document']['idObjType'] == 6:
        if hit['Design'].get('DocFlow'):
            hit['Design']['DocFlow']['Documents'].sort(
                key=lambda document: time.strptime(document['DocRecord'].get(
                    'DocRegDate', '1970-01-01'
                ), "%Y-%m-%d")
            )
        if hit['Design'].get('PaymentDetails'):
            hit['Design']['PaymentDetails']['Payment'].sort(
                key=lambda document: time.strptime(document.get(
                    'PaymentDate', '1970-01-01'
                ), "%Y-%m-%d")
            )


def get_registration_status_color(hit):
    """Возвращает цвте статуса охранного документа объекта пром. собств."""
    status = 'gray'
    if hit['Document']['idObjType'] in (1, 2, 3):
        try:
            if hit['Document']['RegistrationStatus'] == 'A':
                status = 'green'
            elif hit['Document']['RegistrationStatus'] == 'N':
                status = 'red'
            elif hit['Document']['RegistrationStatus'] == 'T':
                status = 'yellow'
            else:
                status = 'red'
        except KeyError:
            status = 'red'
    elif hit['Document']['idObjType'] == 4:
        status = 'green'

        red_transaction_types = [
            'TerminationNoRenewalFee',
            'TotalTerminationByOwner',
            'TotalInvalidationByCourt',
            'TotalTerminationByCourt',
            'TotalInvalidationByAppeal',
        ]

        if hit.get('TradeMark', {}).get('Transactions'):
            last_transaction_type = hit['TradeMark']['Transactions']['Transaction'][
                len(hit['TradeMark']['Transactions']['Transaction']) - 1].get('@type')

            if last_transaction_type in red_transaction_types:
                status = 'red'

    elif hit['Document']['idObjType'] == 5:
        status = 'green'

    elif hit['Document']['idObjType'] == 6:
        status = 'green'

        red_transaction_types = [
            'Termination',
            'TerminationByAppeal',
            'TerminationNoRenewalFee',
            'TotalInvalidationByAppeal',
            'TotalInvalidationByCourt',
            'TotalTerminationByOwner',
        ]

        if hit.get('Design', {}).get('Transactions'):
            last_transaction_type = hit['Design']['Transactions']['Transaction'][
                len(hit['Design']['Transactions']['Transaction']) - 1].get('@type')

            if last_transaction_type in red_transaction_types:
                status = 'red'

    # ТМ (Мадрид)
    elif hit['Document']['idObjType'] in (9, 14):
        status = 'green'

    return status


def get_fixed_mark_status_code(app_data):
    """Анализирует список документов и возвращает код статуса согласно их наличию."""
    result = int(app_data['Document'].get('MarkCurrentStatusCodeType', 0))
    for doc in app_data['TradeMark'].get('DocFlow', {}).get('Documents', []):
        if ('ТM-1.1' in doc['DocRecord']['DocType'] or 'ТМ-1.1' in doc['DocRecord']['DocType']) and result < 2000:
            result = 3000
        if 'Т-05' in doc['DocRecord']['DocType'] and result < 3000:
            result = 3000
        if 'Т-08' in doc['DocRecord']['DocType'] and result < 4000:
            result = 4000
    return result


def filter_app_data(app_data, user):
    """Фильтрует данные заявки. Оставляет только необходимую и доступную для отображения информацию."""

    # Если это заявка на полезную модель или пром образец, заявка на ТМ без установленной даты подачи
    # то необходимо убрать всю "закрытую" информацию
    # (кроме как для вип-пользователей или людей, которые имеют отношение к заявке)
    if app_data['search_data']['obj_state'] == 1 and not user_has_access_to_docs(user, app_data):

        if app_data['Document']['idObjType'] == 1 and not app_data['Claim'].get('I_43.D'):  # Изобретения
            res = {
                'meta': app_data['meta'],
                'Document': app_data['Document'],
                'DOCFLOW': app_data.get('DOCFLOW'),
                'search_data': {
                    'app_number': app_data['search_data']['app_number'],
                    'obj_state': app_data['search_data']['obj_state'],
                }
            }
            return res

        elif app_data['Document']['idObjType'] == 2:  # Полезные модели
            res = {
                'meta': app_data['meta'],
                'Document': app_data['Document'],
                'DOCFLOW': app_data.get('DOCFLOW'),
                'search_data': {
                    'app_number': app_data['search_data']['app_number'],
                    'obj_state': app_data['search_data']['obj_state'],
                }
            }
            return res

        elif app_data['Document']['idObjType'] == 4:
            # mark_status = int(app_data['Document'].get('MarkCurrentStatusCodeType', 0))
            mark_status = get_fixed_mark_status_code(app_data)
            app_date = datetime.datetime.strptime(app_data['search_data']['app_date'][:10], '%Y-%m-%d')
            if app_data['TradeMark']['TrademarkDetails'].get('Code_441'):
                date_441 = datetime.datetime.strptime(
                    app_data['TradeMark']['TrademarkDetails'].get('Code_441'),
                    '%Y-%m-%d'
                )
            else:
                date_441 = None
            date_441_start = datetime.datetime.strptime('2020-08-18', '%Y-%m-%d')

            # Условие, которое определяет установлена ли дата подачи заявки
            if (mark_status < 2000 and app_date < date_441_start) or (date_441 is None and app_date >= date_441_start):
                res = {}
                res.update({'meta': app_data['meta']})
                res.update({'Document': app_data['Document']})
                res.update(
                    {
                        'TradeMark': {
                            'TrademarkDetails': {
                                'stages': app_data['TradeMark'].get('TrademarkDetails', {}).get('stages', []),
                                'application_status': app_data['TradeMark'].get('TrademarkDetails', {}).get(
                                    'application_status'
                                )
                            },
                            'DocFlow': app_data['TradeMark'].get('DocFlow', {}),
                        }
                    }
                )
                res.update({'search_data': {
                    'app_number': app_data['search_data']['app_number'],
                    'obj_state': app_data['search_data']['obj_state'],
                }})
                return res

        elif app_data['Document']['idObjType'] == 6:  # Пром. образцы
            res = {
                'meta': app_data['meta'],
                'Document': app_data['Document'],
                'Design': {
                    'DesignDetails': {
                        'stages': app_data['Design'].get('DesignDetails', {}).get('stages', []),
                        'application_status': app_data['Design'].get('DesignDetails', {}).get(
                            'application_status'
                        )
                    },
                    'DocFlow': app_data['Design'].get('DocFlow')
                },
                'search_data': {
                    'app_number': app_data['search_data']['app_number'],
                    'obj_state': app_data['search_data']['obj_state'],
                }
            }
            return res

    return app_data


def is_app_limited(app_data, user):
    """Является ли заявка такой (для пользователя), библиографические данные которой не должны публиковаться"""
    if app_data['search_data']['obj_state'] == 1 and not user_has_access_to_docs(user, app_data):
        if app_data['Document']['idObjType'] == 1 and not app_data['Claim'].get('I_43.D'):  # Изобретения
            return True
        elif app_data['Document']['idObjType'] == 2:  # Полезные модели
            return True
        elif app_data['Document']['idObjType'] == 4:  # Заявки на ТМ
            # mark_status = int(app_data['Document'].get('MarkCurrentStatusCodeType', 0))
            mark_status = get_fixed_mark_status_code(app_data)
            app_date = datetime.datetime.strptime(app_data['search_data']['app_date'][:10], '%Y-%m-%d')
            if app_data['TradeMark']['TrademarkDetails'].get('Code_441'):
                date_441 = datetime.datetime.strptime(
                    app_data['TradeMark']['TrademarkDetails'].get('Code_441'),
                    '%Y-%m-%d'
                )
            else:
                date_441 = None
            date_441_start = datetime.datetime.strptime('2020-08-18', '%Y-%m-%d')

            # Условие, которое определяет установлена ли дата подачи заявки
            return (mark_status < 2000 and app_date < date_441_start) or (date_441 is None and app_date >= date_441_start)
        elif app_data['Document']['idObjType'] == 6:  # Пром. образцы
            return True
    return False


def get_ipc_codes_with_schedules(lang_code):
    """Возвращает ИНИД коды с реестрами, в оторый они входят"""
    qs = IpcCode.objects.prefetch_related(
        'obj_types',
        Prefetch(
            'inidcodeschedule_set',
            queryset=InidCodeSchedule.objects.filter(
                enable_search=True
            ).exclude(
                schedule_type__id__gt=19,
                schedule_type__id__lt=30
            ).prefetch_related('schedule_type')
        ),
    ).exclude(
        obj_types=None,
    ).exclude(
        inidcodeschedule__schedule_type__id__gt=19,
        inidcodeschedule__schedule_type__id__lt=30,
        inidcodeschedule__schedule_type__ipccode__isnull=True,
    )

    res = []
    for item in qs:
        obj_states = [
            2 if schedule_type.schedule_type.id in (3, 4, 5, 6, 7, 8, 16, 17, 18, 19, 30, 32) else 1
            for schedule_type in item.inidcodeschedule_set.all()
        ]

        if obj_states:
            res.append({
                'id': item.id,
                'value': item.code_value_ua if lang_code == 'ua' else item.code_value_en,
                'data_type': item.code_data_type,
                'obj_types': [obj_type.id for obj_type in item.obj_types.all()],
                'obj_states': obj_states,
            })
    return res
